#!/usr/bin/env python3
"""
CIBE 美业数据白皮书自动生成系统 — 后端代理
功能：文件解析、DeepSeek 内容生成、豆包星绘配图生成
"""

import os
import io
import json
import time
import base64
import re
import html
import hashlib
import hmac
import asyncio
import contextvars
import uuid
import sqlite3
import queue
from abc import ABC, abstractmethod
from typing import Any, Optional, List, Dict, Tuple
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

import pandas as pd
import pdfplumber
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
import requests
import uvicorn

# Import reflection and web search modules
from reflection_agent import ReflectionAgent, ReflectionSession
from reflection_criteria import ReflectionCriteria
from reflection_storage import DEFAULT_REFLECTION_STORAGE_PATH, ReflectionStorage
from web_search import WebSearch
from web_access_research import (
    collect_recent_references,
    collect_recent_references_debug,
    build_research_context,
)

# ---------------------------------------------------------------------------
# 配置
# ---------------------------------------------------------------------------
load_dotenv()

DEEPSEEK_API_KEY=os.getenv("DEEPSEEK_API_KEY", "")
ARK_API_KEY=os.getenv("ARK_API_KEY", "")
MODEL_NAME=os.getenv("MODEL_NAME", "deepseek-chat")
ALLOWED_ORIGINS = [
    o.strip()
    for o in os.getenv("ALLOWED_ORIGINS", "http://localhost:5678,http://127.0.0.1:5678").split(",")
]

# Reflection and verification configuration
ENABLE_REFLECTION = os.getenv("ENABLE_REFLECTION", "true").lower() == "true"
ENABLE_WEB_SEARCH = os.getenv("ENABLE_WEB_SEARCH", "true").lower() == "true"
ENABLE_DATA_VERIFICATION = os.getenv("ENABLE_DATA_VERIFICATION", "false").lower() == "true"
REFLECTION_STORAGE_PATH = os.getenv("REFLECTION_STORAGE_PATH", DEFAULT_REFLECTION_STORAGE_PATH)
WEB_SEARCH_TIMEOUT = int(os.getenv("WEB_SEARCH_TIMEOUT", "10"))
PIPELINE_MEMORY_PATH = os.getenv("PIPELINE_MEMORY_PATH", os.path.join("data", "pipeline_memory.json"))
FAST_REFLECTION_MIN_CHARS = int(os.getenv("FAST_REFLECTION_MIN_CHARS", "8000"))
FAST_REFLECTION_MIN_REFS = int(os.getenv("FAST_REFLECTION_MIN_REFS", "6"))
IMAGE_PROMPT_CACHE_PATH = os.getenv("IMAGE_PROMPT_CACHE_PATH", os.path.join("data", "image_prompt_cache.json"))
IMAGE_PROMPT_CACHE_TTL = int(os.getenv("IMAGE_PROMPT_CACHE_TTL", "604800"))
AUTH_ENABLED = os.getenv("AUTH_ENABLED", "true").lower() == "true"
AUTH_USERNAME = os.getenv("AUTH_USERNAME", "admin")
AUTH_PASSWORD = os.getenv("AUTH_PASSWORD", "CIBE@2026")
AUTH_COOKIE_NAME = os.getenv("AUTH_COOKIE_NAME", "cibe_session")
AUTH_SESSION_TTL = int(os.getenv("AUTH_SESSION_TTL", "86400"))
AUTH_SECRET = os.getenv("AUTH_SECRET") or hashlib.sha256(
    f"{DEEPSEEK_API_KEY}:{AUTH_PASSWORD}:cibe-auth".encode("utf-8")
).hexdigest()
APP_DB_PATH = os.getenv("APP_DB_PATH", os.path.join("data", "app.db"))
GENERATION_WORKERS = int(os.getenv("GENERATION_WORKERS", "2"))

# ---------------------------------------------------------------------------
# FastAPI 应用
# ---------------------------------------------------------------------------
app = FastAPI(title="CIBE 白皮书生成器")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _auth_signature(payload: str) -> str:
    return hmac.new(AUTH_SECRET.encode("utf-8"), payload.encode("utf-8"), hashlib.sha256).hexdigest()


def _create_session_token(username: str) -> str:
    expires = int(time.time()) + AUTH_SESSION_TTL
    nonce = uuid.uuid4().hex
    payload = f"{username}|{expires}|{nonce}"
    return f"{payload}|{_auth_signature(payload)}"


def _verify_session_token(token: str) -> Optional[str]:
    try:
        username, expires_raw, nonce, signature = (token or "").split("|", 3)
        payload = f"{username}|{expires_raw}|{nonce}"
        if not hmac.compare_digest(signature, _auth_signature(payload)):
            return None
        if int(expires_raw) < int(time.time()):
            return None
        if not _user_exists(username):
            return None
        return username
    except Exception:
        return None


def _is_authenticated(request: Request) -> bool:
    if not AUTH_ENABLED:
        return True
    return bool(_verify_session_token(request.cookies.get(AUTH_COOKIE_NAME, "")))


def _auth_required_response() -> JSONResponse:
    return JSONResponse(
        {"status": "unauthorized", "detail": "请先登录"},
        status_code=401,
    )


def _db_connect() -> sqlite3.Connection:
    os.makedirs(os.path.dirname(APP_DB_PATH) or ".", exist_ok=True)
    conn = sqlite3.connect(APP_DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def _hash_password(password: str, salt: Optional[str] = None) -> str:
    salt = salt or base64.urlsafe_b64encode(os.urandom(16)).decode("ascii")
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 120_000)
    return f"pbkdf2_sha256${salt}${base64.b64encode(digest).decode('ascii')}"


def _verify_password(password: str, stored_hash: str) -> bool:
    try:
        algorithm, salt, digest = stored_hash.split("$", 2)
        if algorithm != "pbkdf2_sha256":
            return False
        expected = _hash_password(password, salt).split("$", 2)[2]
        return hmac.compare_digest(expected, digest)
    except Exception:
        return False


def _init_app_db() -> None:
    with _db_connect() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL DEFAULT 'admin',
                created_at INTEGER NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS generation_jobs (
                id TEXT PRIMARY KEY,
                username TEXT NOT NULL,
                status TEXT NOT NULL,
                data_text TEXT NOT NULL,
                markdown TEXT,
                error TEXT,
                created_at INTEGER NOT NULL,
                started_at INTEGER,
                finished_at INTEGER
            )
            """
        )
        user_count = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        if user_count == 0:
            conn.execute(
                "INSERT INTO users(username, password_hash, role, created_at) VALUES (?, ?, ?, ?)",
                (AUTH_USERNAME, _hash_password(AUTH_PASSWORD), "admin", int(time.time())),
            )


def _verify_user_password(username: str, password: str) -> bool:
    with _db_connect() as conn:
        row = conn.execute("SELECT password_hash FROM users WHERE username=?", (username,)).fetchone()
    return bool(row and _verify_password(password, row["password_hash"]))


def _get_user(username: str) -> Optional[Dict[str, Any]]:
    with _db_connect() as conn:
        row = conn.execute(
            "SELECT id, username, role, created_at FROM users WHERE username=?",
            (username,),
        ).fetchone()
    return _row_to_job(row) if row else None


def _list_users(limit: int = 100) -> List[Dict[str, Any]]:
    with _db_connect() as conn:
        rows = conn.execute(
            """
            SELECT id, username, role, created_at
            FROM users
            ORDER BY created_at ASC, id ASC
            LIMIT ?
            """,
            (limit,),
        ).fetchall()
    return [_row_to_job(row) for row in rows]


def _count_users() -> int:
    with _db_connect() as conn:
        return int(conn.execute("SELECT COUNT(*) FROM users").fetchone()[0])


def _count_generation_jobs() -> int:
    with _db_connect() as conn:
        return int(conn.execute("SELECT COUNT(*) FROM generation_jobs").fetchone()[0])


def _create_user(username: str, password: str, role: str = "user") -> Dict[str, Any]:
    cleaned_username = (username or "").strip()
    cleaned_role = (role or "user").strip().lower()
    if not cleaned_username:
        raise ValueError("账号不能为空")
    if len(cleaned_username) < 3:
        raise ValueError("账号至少 3 个字符")
    if not re.fullmatch(r"[A-Za-z0-9_.-]+", cleaned_username):
        raise ValueError("账号仅支持字母、数字、点、横线和下划线")
    if len(password or "") < 6:
        raise ValueError("密码至少 6 位")
    if cleaned_role not in {"admin", "user"}:
        raise ValueError("角色不合法")

    now = int(time.time())
    with _db_connect() as conn:
        conn.execute(
            """
            INSERT INTO users(username, password_hash, role, created_at)
            VALUES (?, ?, ?, ?)
            """,
            (cleaned_username, _hash_password(password), cleaned_role, now),
        )
    return {
        "username": cleaned_username,
        "password": password,
        "role": cleaned_role,
        "created_at": now,
    }


def _generate_password(length: int = 10) -> str:
    alphabet = "ABCDEFGHJKLMNPQRSTUVWXYZabcdefghijkmnopqrstuvwxyz23456789"
    return "".join(alphabet[ord(os.urandom(1)) % len(alphabet)] for _ in range(length))


def _batch_create_users(prefix: str = "tester", count: int = 5, role: str = "user") -> List[Dict[str, Any]]:
    normalized_prefix = (prefix or "tester").strip().lower()
    if not re.fullmatch(r"[A-Za-z0-9_.-]+", normalized_prefix):
        raise ValueError("账号前缀不合法")
    if count < 1 or count > 50:
        raise ValueError("批量创建数量需在 1-50 之间")

    created: List[Dict[str, Any]] = []
    next_index = 1
    while len(created) < count:
        username = f"{normalized_prefix}{next_index:02d}"
        next_index += 1
        if _user_exists(username):
            continue
        password = _generate_password()
        created.append(_create_user(username, password, role=role))
    return created


def _user_exists(username: str) -> bool:
    try:
        with _db_connect() as conn:
            row = conn.execute("SELECT 1 FROM users WHERE username=?", (username,)).fetchone()
        return bool(row)
    except Exception:
        # During early startup or tests before DB init, keep the env admin usable.
        return username == AUTH_USERNAME


def _create_generation_job(username: str, data_text: str, job_id: Optional[str] = None) -> str:
    new_job_id = (job_id or f"job-{uuid.uuid4().hex}").strip()
    now = int(time.time())
    with _db_connect() as conn:
        conn.execute(
            """
            INSERT INTO generation_jobs(id, username, status, data_text, created_at)
            VALUES (?, ?, 'queued', ?, ?)
            """,
            (new_job_id, username, data_text, now),
        )
    _emit_progress("[Queue] 任务已提交，等待后台生成", job_id=new_job_id, stage="queue")
    return new_job_id


def _row_to_job(row: sqlite3.Row) -> Dict[str, Any]:
    return dict(row)


def _get_generation_job(job_id: str) -> Optional[Dict[str, Any]]:
    with _db_connect() as conn:
        row = conn.execute("SELECT * FROM generation_jobs WHERE id=?", (job_id,)).fetchone()
    return _row_to_job(row) if row else None


def _list_generation_jobs(username: str, limit: int = 20) -> List[Dict[str, Any]]:
    with _db_connect() as conn:
        rows = conn.execute(
            """
            SELECT id, username, status, error, created_at, started_at, finished_at
            FROM generation_jobs
            WHERE username=?
            ORDER BY created_at DESC
            LIMIT ?
            """,
            (username, limit),
        ).fetchall()
    return [_row_to_job(row) for row in rows]


def _list_all_generation_jobs(limit: int = 50) -> List[Dict[str, Any]]:
    with _db_connect() as conn:
        rows = conn.execute(
            """
            SELECT id, username, status, error, created_at, started_at, finished_at, markdown, data_text
            FROM generation_jobs
            ORDER BY created_at DESC
            LIMIT ?
            """,
            (limit,),
        ).fetchall()
    return [_row_to_job(row) for row in rows]


def _extract_job_title(job: Dict[str, Any]) -> str:
    markdown = job.get("markdown") or ""
    match = re.search(r"^#\s+(.+)$", markdown, flags=re.MULTILINE)
    if match:
        return match.group(1).strip()
    raw_text = job.get("data_text") or ""
    raw_text = re.sub(r"【文件:\s*[^】]+】", " ", raw_text)
    raw_text = raw_text.replace("【文字输入】", " ")
    lines = [line.strip() for line in raw_text.splitlines()]
    lines = [line for line in lines if line and line != "---"]
    for line in lines:
        normalized = re.sub(r"\s+", " ", line).strip(" -|")
        if not normalized:
            continue
        if "：" in normalized:
            prefix, value = normalized.split("：", 1)
            if prefix.strip() in {"品牌名称", "品牌", "产品", "关键词", "主题", "输入"} and value.strip():
                return value.strip()[:48]
        if ":" in normalized:
            prefix, value = normalized.split(":", 1)
            if prefix.strip() in {"brand", "product", "topic", "keyword"} and value.strip():
                return value.strip()[:48]
        return normalized[:48]
    return job.get("id", "未命名任务")


def _build_admin_dashboard(limit: int = 30) -> Dict[str, Any]:
    jobs = _list_all_generation_jobs(limit=limit)
    return {
        "metrics": {
            "user_count": _count_users(),
            "job_count": _count_generation_jobs(),
        },
        "recent_jobs": [
            {
                "id": job["id"],
                "username": job["username"],
                "status": job["status"],
                "error": job.get("error"),
                "created_at": job.get("created_at"),
                "started_at": job.get("started_at"),
                "finished_at": job.get("finished_at"),
                "title": _extract_job_title(job),
                "preview": ((job.get("markdown") or job.get("data_text") or "")[:180]).strip(),
            }
            for job in jobs
        ],
    }


def _update_generation_job(job_id: str, **fields: Any) -> None:
    if not fields:
        return
    columns = ", ".join(f"{key}=?" for key in fields)
    values = list(fields.values()) + [job_id]
    with _db_connect() as conn:
        conn.execute(f"UPDATE generation_jobs SET {columns} WHERE id=?", values)


def _delete_generation_job(job_id: str) -> None:
    with _db_connect() as conn:
        conn.execute("DELETE FROM generation_jobs WHERE id=?", (job_id,))


_GENERATION_QUEUE: "queue.Queue[str]" = queue.Queue()
_WORKERS_STARTED = False


def _enqueue_generation_job(job_id: str) -> None:
    _GENERATION_QUEUE.put(job_id)


def _run_generation_job(job_id: str) -> None:
    job = _get_generation_job(job_id)
    if not job:
        return
    token = _CURRENT_PROGRESS_JOB.set(job_id)
    try:
        _update_generation_job(job_id, status="running", started_at=int(time.time()))
        _emit_progress("[Queue] 后台 worker 已开始处理任务", job_id=job_id, stage="queue")
        markdown = _generate_whitepaper_sync(job["data_text"])
        _update_generation_job(
            job_id,
            status="succeeded",
            markdown=markdown,
            finished_at=int(time.time()),
        )
        _finish_progress(job_id, "白皮书生成完成")
    except Exception as exc:
        _update_generation_job(
            job_id,
            status="failed",
            error=str(exc),
            finished_at=int(time.time()),
        )
        _finish_progress(job_id, f"生成失败: {exc}")
    finally:
        _CURRENT_PROGRESS_JOB.reset(token)


def _generation_worker_loop() -> None:
    while True:
        job_id = _GENERATION_QUEUE.get()
        try:
            _run_generation_job(job_id)
        finally:
            _GENERATION_QUEUE.task_done()


def _start_generation_workers() -> None:
    global _WORKERS_STARTED
    if _WORKERS_STARTED:
        return
    _WORKERS_STARTED = True
    for idx in range(max(1, GENERATION_WORKERS)):
        worker = threading.Thread(
            target=_generation_worker_loop,
            name=f"generation-worker-{idx + 1}",
            daemon=True,
        )
        worker.start()


@app.middleware("http")
async def auth_middleware(request: Request, call_next):
    path = request.url.path
    public_paths = {
        "/",
        "/favicon.ico",
        "/api/health",
        "/api/auth/login",
        "/api/auth/logout",
        "/api/auth/me",
    }
    if (
        AUTH_ENABLED
        and path.startswith("/api/")
        and path not in public_paths
        and not _is_authenticated(request)
    ):
        return _auth_required_response()
    return await call_next(request)


# 静态文件：让 index.html 可以直接通过 / 访问
app.mount("/static", StaticFiles(directory="."), name="static")


# ---------------------------------------------------------------------------
# 文件解析引擎
# ---------------------------------------------------------------------------
class FileParser:
    """支持 CSV / Excel(xlsx/xls) / Word(docx) / PDF 四类格式解析"""

    @staticmethod
    def parse(filename: str, content: bytes) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        parsers = {
            "csv": FileParser._parse_csv,
            "xlsx": FileParser._parse_csv,
            "xls": FileParser._parse_csv,
            "docx": FileParser._parse_docx,
            "pdf": FileParser._parse_pdf,
        }
        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"不支持的文件格式: .{ext}，请上传 CSV / Word / PDF 文件")
        return parser(content, filename)

    @staticmethod
    def _parse_csv(content: bytes, filename: str = "") -> str:
        """CSV/Excel → 结构化摘要"""
        try:
            ext = filename.rsplit(".", 1)[-1].lower() if filename else "csv"
            if ext in ("xlsx", "xls"):
                df = pd.read_excel(io.BytesIO(content))
            else:
                # 尝试多种编码
                for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                    try:
                        df = pd.read_csv(io.BytesIO(content), encoding=enc)
                        break
                    except (UnicodeDecodeError, Exception):
                        continue
                else:
                    df = pd.read_csv(io.BytesIO(content), encoding="utf-8", errors="replace")

            lines = [
                f"【数据文件摘要】",
                f"行数: {len(df)}，列数: {len(df.columns)}",
                f"列名: {', '.join(df.columns.tolist())}",
                "",
                "【列类型与统计】",
            ]
            for col in df.columns:
                dtype = str(df[col].dtype)
                if pd.api.types.is_numeric_dtype(df[col]):
                    desc = df[col].describe()
                    lines.append(
                        f"  {col} ({dtype}): 均值={desc['mean']:.2f}, "
                        f"最小={desc['min']:.2f}, 最大={desc['max']:.2f}, "
                        f"中位数={desc['50%']:.2f}"
                    )
                else:
                    nunique = df[col].nunique()
                    top_vals = df[col].value_counts().head(5).to_dict()
                    lines.append(f"  {col} ({dtype}): {nunique}个唯一值, Top5: {top_vals}")

            lines.append("")
            lines.append("【前5行样本数据】")
            lines.append(df.head(5).to_string(index=False))
            return "\n".join(lines)
        except Exception as e:
            raise ValueError(f"CSV 解析失败: {str(e)}")

    @staticmethod
    def _parse_docx(content: bytes, filename: str = "") -> str:
        """Word → 纯文本"""
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                raise ValueError("Word 文档内容为空")
            return "\n".join(paragraphs)
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Word 解析失败: {str(e)}")

    @staticmethod
    def _parse_pdf(content: bytes, filename: str = "") -> str:
        """PDF → 纯文本"""
        try:
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            if not text_parts:
                raise ValueError("PDF 文档未提取到文字内容")
            return "\n\n".join(text_parts)
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"PDF 解析失败: {str(e)}")


# ---------------------------------------------------------------------------
# LLM 抽象层
# ---------------------------------------------------------------------------
class LLMProvider(ABC):
    @abstractmethod
    def generate(self, prompt: str, system: str) -> str:
        ...


class DeepSeekProvider(LLMProvider):
    # 全局信号量：限制同时最多 3 个 API 并发请求
    _semaphore = threading.Semaphore(3)

    def __init__(self):
        if not DEEPSEEK_API_KEY:
            raise RuntimeError("DEEPSEEK_API_KEY 未配置，请检查 .env 文件")
        self.client = OpenAI(
            api_key=DEEPSEEK_API_KEY,
            base_url="https://api.deepseek.com",
        )

    def generate(self, prompt: str, system: str, max_retries: int = 3,
                 max_tokens: int = 8192) -> str:
        """带并发控制和重试的 API 调用"""
        with self._semaphore:
            for attempt in range(max_retries):
                try:
                    response = self.client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[
                            {"role": "system", "content": system},
                            {"role": "user", "content": prompt},
                        ],
                        temperature=0.7,
                        max_tokens=max_tokens,
                    )
                    return response.choices[0].message.content
                except Exception as e:
                    wait = 2 ** attempt  # 1s, 2s, 4s 指数退避
                    print(f"[DeepSeek] 请求失败(第{attempt+1}次): {e}, {wait}s后重试...")
                    if attempt == max_retries - 1:
                        raise
                    time.sleep(wait)


# 预留：未来可切换其他 Provider
# class GeminiProvider(LLMProvider): ...
# class GPT4oProvider(LLMProvider): ...


def get_llm_provider() -> LLMProvider:
    """工厂方法：当前返回 DeepSeek，未来可按配置切换"""
    return DeepSeekProvider()


KEYWORD_EXTRACT_SYSTEM = """你是一个中文美妆行业检索关键词拆分器。只输出 JSON，不要解释。"""
KEYWORD_EXTRACT_USER = """从用户输入中提取用于联网检索的关键词。
要求：
1. brand_terms：品牌名，例如「雅诗兰黛」「谷雨」「自然堂」。
2. product_terms：产品/品类词，例如「眼霜」「面膜」「修复精华」。
3. compound_terms：品牌+产品组合原词，例如「雅诗兰黛眼霜」「谷雨面膜」。
4. 不要输出「文字输入」「报告」「数据」「分析」这类系统词或泛词。
5. 如果只有一个连续中文词，也必须尝试拆成品牌和产品。

用户输入：
{text}

输出 JSON Schema：
{{
  "brand_terms": ["品牌"],
  "product_terms": ["产品"],
  "compound_terms": ["品牌产品"]
}}"""


def _fallback_research_topic(data_text: str) -> str:
    text = re.sub(r"【[^【】]{0,20}】", " ", data_text or "")
    text = re.sub(r"\s+", " ", text).strip()
    return text[:120]


def _extract_research_topic_with_ai(data_text: str) -> str:
    fallback = _fallback_research_topic(data_text)
    if not fallback:
        return ""
    try:
        llm = get_llm_provider()
        raw = llm.generate(
            prompt=KEYWORD_EXTRACT_USER.format(text=fallback[:500]),
            system=KEYWORD_EXTRACT_SYSTEM,
        )
        parsed = _parse_llm_json(raw)
        terms: List[str] = []
        for key in ("brand_terms", "product_terms", "compound_terms"):
            values = parsed.get(key, [])
            if isinstance(values, str):
                values = [values]
            if isinstance(values, list):
                for value in values:
                    term = re.sub(r"\s+", "", str(value or "").strip())
                    if term and term not in {"文字输入", "报告", "数据", "分析", "市场"}:
                        terms.append(term)
        terms = list(dict.fromkeys(terms))
        if terms:
            search_topic = " ".join(terms[:6])
            print(f"[Research][Topic] AI拆词: {search_topic}")
            return search_topic
    except Exception as exc:
        print(f"[Research][Topic] AI拆词失败，回退原输入: {exc}")
    return fallback


# ---------------------------------------------------------------------------
# 豆包星绘 — 配图生成
# ---------------------------------------------------------------------------
class ImageGenerator:
    """
    豆包星绘 — 通过火山方舟 (Volcengine ARK) OpenAI 兼容接口生成图片
    文档: https://www.volcengine.com/docs/82379/1399424
    """

    ENDPOINT = "https://ark.cn-beijing.volces.com/api/v3/images/generations"

    def __init__(self):
        self.api_key = ARK_API_KEY

    @property
    def available(self) -> bool:
        return bool(self.api_key)

    def generate(self, prompt: str, width: int = 768, height: int = 512) -> Optional[str]:
        """
        生成图片，返回 base64 编码的图片数据。
        如果未配置 Key 或生成失败，返回 None。
        """
        if not self.available:
            return None

        try:
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}",
            }

            body = {
                "model": "doubao-seedream-3-0-t2i-250415",
                "prompt": prompt,
                "size": f"{width}x{height}",
                "response_format": "b64_json",
                "watermark": False,
                "n": 1,
            }

            resp = requests.post(
                self.ENDPOINT,
                headers=headers,
                json=body,
                timeout=90,
            )
            resp.raise_for_status()
            result = resp.json()

            # OpenAI 兼容格式: {"data": [{"b64_json": "..."}]}
            data_list = result.get("data", [])
            if data_list and data_list[0].get("b64_json"):
                return data_list[0]["b64_json"]

            # 也兼容 url 模式
            if data_list and data_list[0].get("url"):
                img_resp = requests.get(data_list[0]["url"], timeout=30)
                return base64.b64encode(img_resp.content).decode("utf-8")

            print(f"[星绘] 返回数据无图片: {json.dumps(result, ensure_ascii=False)[:200]}")
            return None

        except requests.exceptions.HTTPError as e:
            err_body = ""
            try:
                err_body = e.response.text[:300]
            except Exception:
                pass
            print(f"[星绘] HTTP 错误 {e.response.status_code}: {err_body}")
            return None
        except Exception as e:
            print(f"[星绘] 请求异常: {e}")
            return None


image_generator = ImageGenerator()


# ---------------------------------------------------------------------------
# 前端流式进度
# ---------------------------------------------------------------------------
_CURRENT_PROGRESS_JOB: contextvars.ContextVar[Optional[str]] = contextvars.ContextVar(
    "current_progress_job",
    default=None,
)
_PROGRESS_LOCK = threading.Lock()
_PROGRESS_EVENTS: Dict[str, List[Dict[str, Any]]] = {}
_PROGRESS_DONE: Dict[str, bool] = {}


def _emit_progress(message: str, job_id: Optional[str] = None, stage: str = "log") -> None:
    target_job = job_id or _CURRENT_PROGRESS_JOB.get()
    if not target_job:
        return
    text = str(message or "").strip()
    if not text:
        return
    with _PROGRESS_LOCK:
        events = _PROGRESS_EVENTS.setdefault(target_job, [])
        events.append(
            {
                "id": len(events) + 1,
                "stage": stage,
                "message": text,
                "timestamp": int(time.time()),
            }
        )
        if len(events) > 500:
            del events[: len(events) - 500]


def _finish_progress(job_id: Optional[str], message: str = "生成完成") -> None:
    if not job_id:
        return
    _emit_progress(message, job_id=job_id, stage="done")
    with _PROGRESS_LOCK:
        _PROGRESS_DONE[job_id] = True


def _progress_print(message: str, stage: str = "log") -> None:
    print(message)
    _emit_progress(message, stage=stage)


@app.on_event("startup")
async def startup_event():
    _init_app_db()
    _start_generation_workers()


# ---------------------------------------------------------------------------
# 反思功能存储
# ---------------------------------------------------------------------------
class ReflectionStorage:
    """反思报告存储管理器"""
    
    def __init__(self):
        self.reflections = []
        self.lock = threading.Lock()
    
    def add_reflection(self, report_id: str, content: str, metadata: dict):
        """添加反思报告"""
        with self.lock:
            reflection = {
                "id": report_id,
                "content": content,
                "metadata": metadata,
                "timestamp": datetime.now().isoformat()
            }
            self.reflections.append(reflection)
            return reflection
    
    def get_reflection(self, report_id: str):
        """获取指定反思报告"""
        with self.lock:
            for reflection in self.reflections:
                if reflection["id"] == report_id:
                    return reflection
            return None
    
    def get_reflection_history(self, limit: int = 50, offset: int = 0):
        """获取反思历史"""
        with self.lock:
            # 按时间倒序排列
            sorted_reflections = sorted(
                self.reflections, 
                key=lambda x: x["timestamp"], 
                reverse=True
            )
            return sorted_reflections[offset:offset + limit]
    
    def get_reflections_by_type(self, reflection_type: str, limit: int = 50):
        """按类型获取反思报告"""
        with self.lock:
            filtered = [
                r for r in self.reflections 
                if r.get("metadata", {}).get("type") == reflection_type
            ]
            # 按时间倒序排列
            return sorted(filtered, key=lambda x: x["timestamp"], reverse=True)[:limit]

# 全局反思存储实例
reflection_storage = ReflectionStorage()

# ---------------------------------------------------------------------------
# 四角色 Prompt 体系（专业报告级）
# ---------------------------------------------------------------------------

# ========================
# 角色一：白皮书战略官
# ========================
AGENT1_SYSTEM = """# Role: 行业白皮书首席架构师

你是一位为 CBNData、艾瑞咨询、KPMG、德勤等机构操盘过数十份行业白皮书的首席架构师。
你的任务是将原始数据设计成一份面向决策层的标准行业白皮书结构化大纲（JSON），后续将据此自动扩写正文。

# 核心方法论

## 1. 洋葱递进
章节必须形成「宏观背景 → 市场全景 → 细分深挖 → 竞争格局 → 模式创新 → 未来展望」的递进纵深。
严禁平行罗列（如"趋势一/趋势二/趋势三"或多个无推进关系的同级分析）。

## 2. 判断句标题
每章标题必须是结论句/判断句，不是话题词。
✗ "功效护肤市场现状"　✓ "功效护肤正在从营销概念走向临床验证"
✗ "渠道分析"　　　　　✓ "线下体验场景正在重新定义专业品牌的转化效率"

## 3. 完整撰稿 brief
每章必须给出：核心问题、导语、2-4 个二级小节、≥2 个数据锚点、≥1 个案例方向、与上章承接、图表意图、结论小结。

## 4. 研究报告体
全文定位为行业研究报告，不是品牌宣传册、公众号长文或 PPT 提纲。标题稳准硬，禁止空泛判断（"未来可期""机遇与挑战并存"）。

# 视觉规则
图表优先于图片（图表是证据，图片仅辅助）。配图仅限真实可拍摄场景（门店/诊疗/产品/案例现场），禁止抽象概念图和纯氛围图。
必须为至少 3 章设置配图：这些章节的 image_intent 必须是具体可拍摄场景，禁止所有章节 image_intent 都为 none。
若全文只有 4 章，通常选择 3 章配图；若全文 5 章，可选择 3-4 章配图，避免每章都强行配大图。

# 红线
1. 禁止平行堆砌章节、禁止缺失执行摘要/研究说明/结论建议等标准模块。
2. 禁止空话套话营销口号，禁止脱离原始数据胡乱发散。
3. 只输出合法 JSON，不含 ```json 标记，不含任何解释文字。章节严格 4~5 章，每章 sub_sections 不超过 3 个，字段值尽量精简。"""

AGENT1_USER = """基于以下原始数据，设计一份标准行业白皮书的结构化大纲 JSON。

<RAW_DATA>
{data}
</RAW_DATA>

输出 JSON Schema：
{{
  "whitepaper_meta": {{
    "title": "白皮书完整标题（结论导向，不超过25字）",
    "subtitle": "副标题"，100字以内",
    "global_tone": "如：深度洞察、数据叙事、决策参考级",
    "target_audience": "目标读者，如：品牌方高层、渠道决策层、投资与战略团队"
  }},
  "document_spec": {{
    "document_type": "标准行业白皮书",
    "narrative_logic": "本书的总体递进逻辑概述",
    "required_sections": ["封面页","执行摘要","目录页","研究说明/数据口径","正文章节","结论与行动建议","参考说明/附录"],
    "visual_rule": "图表优先，图片从属，避免画册化",
    "writing_rule": "偏研究报告体，不写成营销长文"
  }},
  "front_matter": {{
    "executive_summary": {{
      "summary_positioning": "执行摘要的定位",
      "key_findings": ["关键发现1","关键发现2","关键发现3","关键发现4"],
      "core_data_points": ["核心数据结论1","核心数据结论2","核心数据结论3"],
      "executive_judgement": "面向决策层的一句话总判断"
    }},
    "research_note": {{
      "research_object": "研究对象",
      "time_scope": "时间范围",
      "data_sources": ["数据来源类型1","数据来源类型2"],
      "sample_scope": "样本/案例口径",
      "method_boundary": "研究边界与推演边界",
      "term_definition": "关键名词定义（如无可留空）"
    }}
  }},
  "chapters": [
    {{
      "chapter_id": "01",
      "chapter_title": "判断句式章节标题",
      "core_proposition": "本章要回答的核心问题",
      "chapter_intro": "本章导语",
      "sub_sections": [
        {{"section_id":"01-1","section_title":"二级小节标题","section_focus":"该小节聚焦的问题"}},
        {{"section_id":"01-2","section_title":"二级小节标题","section_focus":"该小节聚焦的问题"}}
      ],
      "data_anchors": ["数据锚点1","数据锚点2"],
      "case_direction": "案例方向（具体到品牌/品类/模式/场景）",
      "content_guidelines": "撰稿brief：核心论点+关键关系+数据引用+案例嵌入+与上章承接",
      "transition_from_previous": "与上一章的逻辑承接关系",
      "needs_chart": true,
      "chart_intent": "图表意图：具体数据关系描述",
      "visual_type": "chart / case_photo / none",
      "image_intent": "必须写具体可拍摄场景；全文至少3章如此填写，其余章节写none",
      "chapter_summary": "本章结论小结"
    }}
  ],
  "back_matter": {{
    "conclusion_summary": {{
      "final_judgement": "全书总结性判断",
      "industry_implications": ["对行业的启示1","对行业的启示2"]
    }},
    "action_recommendations": {{
      "for_brands": ["给品牌方的建议1","给品牌方的建议2"],
      "for_channels": ["给渠道方的建议1","给渠道方的建议2"],
      "for_investment_or_strategy": ["给投资/战略团队的建议1","给投资/战略团队的建议2"]
    }},
    "risk_and_boundary_notes": ["风险提示1","边界提示2"],
    "references_or_appendix": {{
      "required": true,
      "content_direction": "参考数据来源说明、附录指标口径、补充案例说明等"
    }}
  }}
}}

仅输出 JSON："""

# ========================
# 角色二：美业深度主笔
# ========================
AGENT2_SYSTEM = """# Role: 行业白皮书资深撰稿人

你是一位曾在 CBNData、艾瑞研究院担任首席分析师的行业深度写作者。
你的写作风格对标的是《第一财经》深度商业报道和麦肯锡行业洞察报告——
不是在"介绍信息"，而是在"构建论证"。

# 写作方法论（决定你输出质量的核心规则）：

## 1. 叙事结构：总-分-总，论证驱动
每章以一个引人入胜的现象/数据/判断开篇（hook），
中间展开 2-3 个支撑论点（每个论点自成一个完整论证段落），
段落之间用逻辑过渡句自然衔接（因果、转折、递进，而非机械编号）。

## 2. 数据嵌入法：数据是证据，不是装饰
- 每个数据必须服务于一个论点，不能为了"显得专业"而堆砌
- 数据必须有上下文：对比基准（同比/环比/对标谁）、时间跨度、"so what"解读
- 格式范例：
  ✗ "市场规模达到5000亿元。"（孤立数据，无意义）
  ✓ "中国功效护肤品市场在2024年突破5000亿元大关，较三年前几乎翻番。这一增速远超同期大盘个位数增长，折射出消费者正在从'感性种草'向'理性选择'的根本性转变。"

## 3. 禁止清单式写作（最重要的红线）
- 严禁出现"首先...其次...再次...最后..."的机械排列
- 严禁通篇无序列表/有序列表代替段落论述
- 允许在必要时使用极短列表（3项以内）辅助说明，但列表前后必须有段落论述包裹
- 每章至少 80% 内容必须是完整的叙述段落

## 4. 语言风格：专业而不呆板
- 避免AI感重的词汇：「赋能」「助力」「打造」「引领」「新质生产力」等空洞套话
- 用具体动词替代万能动词：不说"推动增长"，说"提高复购、放大客单价或改善渠道周转"
- 允许使用生动的商业比喻和类比，但不过度文学化
- 敢于给出有锐度的判断，不做"正确的废话"
- 严禁输出 XXX、XX%、X亿元、某某 等占位符；没有可信数据时写「公开资料未披露具体数值」，不要编数字。

## 5. 篇幅与密度
- 每章正文 1000-1500 字（不含标题和占位符标签）
- 每 200 字尽量出现 1 个可核验数据点或案例细节；资料不足时用定性判断，不得用占位符凑数
- 禁止开场白（"在当今时代..."）和总结段（"综上所述..."），直接从第一个论点切入

## 6. 标题编号格式（必须严格遵守）
- 一级章节标题格式：## 一、标题内容 / ## 二、标题内容（中文数字编号由系统传入）
- 二级小节标题格式：### 1. 标题内容 / ### 2. 标题内容
- 三级子标题格式：#### （1）标题内容 / #### （2）标题内容
- 各级标题必须单独成行，不与正文混排

## 7. 段落格式
- 正文每一自然段必须是完整书面段落表达
- 严禁使用项目符号（- / * / •）代替主体论述段落
- 允许的例外：极短辅助说明（3项以内）可使用列表，但前后必须有段落论述

## 8. 占位符规范
- 图表占位：在最能支撑核心数据论证的位置插入，格式 `[CHART_PLACEHOLDER]`
- 配图占位：在场景描述或案例论述后插入，格式 `[IMAGE_PLACEHOLDER]`

## 9. 绝对禁止（违反则输出作废）
- 禁止输出任何括号包裹的建议、备注、说明性文字，如：
  ✗ （图表建议：可展示...）
  ✗ （配图建议：...）
  ✗ （注：...）
  ✗ *（建议...）*
- 禁止输出写作自检清单、TODO、编辑批注
- 禁止对图表/配图的内容做文字说明或建议，占位符标签本身就是全部，不需要任何补充文字
- 输出必须是纯粹的、可直接发表的正文，不含任何指导性或元层级的文字"""

AGENT2_USER = """白皮书标题：《{title}》
语调基调：{tone}
目标读者：{target_audience}

请撰写第 {chapter_id} 章：【{chapter_title}】

核心命题：{core_proposition}
本章导语方向：{chapter_intro}
与上一章的承接：{transition}
撰稿 brief：{guidelines}

二级小节结构：
{sub_sections_text}

数据锚点（必须在正文中引用）：
{data_anchors_text}

案例方向：{case_direction}

本章需总结为：{chapter_summary}

图表占位：{chart_placeholder_instruction}
配图占位：{image_placeholder_instruction}

写作检查清单（请在输出前自检，但不要输出清单本身）：
- [ ] 开篇是一个有冲击力的现象/数据/判断，而非"随着...的发展"
- [ ] 没有连续超过3个无序列表项
- [ ] 每个数据都有对比基准和解读
- [ ] 段落之间有逻辑过渡而非机械编号
- [ ] 没有使用「赋能」「助力」「打造」等空洞词汇
- [ ] 字数达到 1000 字以上
- [ ] 没有任何括号建议/备注/说明性文字（如"图表建议：..."、"配图建议：..."、"注：..."）
- [ ] 输出是纯正文，不含任何元层级评论
- [ ] 二级小节按照给定结构展开，使用 ### 1. / ### 2. 格式
- [ ] 章节末尾有简短结论，自然引出下一章

输出格式要求：
- 第一行必须是：## {chapter_number}、{chapter_title}
- 二级小节按给定结构展开，标题格式：### 1. 小节标题 / ### 2. 小节标题
- 三级子标题格式：#### （1）子标题
- 章节末尾用一小段结论收束（不超过3句，不用"综上所述"开头）
- 只输出正文，不要输出自检清单或任何元注释

请立即输出："""

# ========================
# 角色三：视觉叙事导演
# ========================
AGENT3_SYSTEM = """# Role: 高端商业视觉导演

你深度理解「豆包·星绘」文生图模型的底层逻辑，
专注于将商业白皮书中的抽象场景意图，
转化为能被模型精准还原的、具备电影质感的真实摄影指令。
你的每一条 Prompt，都像在向一位顶级商业摄影师下达拍摄简报。

# 豆包·星绘 强制约束红线:
1. 绝对写实：禁止「概念图」「插画」「矢量图」「UI界面」「图表」「几何色块」「赛博朋克」。
   画面必须看起来像一张真实拍摄的高端商业摄影作品。
2. 场景具象化：遇到抽象概念必须落地为可拍摄的真实场景：
   - 「市场增长」→ 「高档美容院奢华接待台特写，玫瑰金装饰细节」
   - 「技术创新」→ 「实验室白大褂研究员手持透明精华液瓶，显微镜背景虚化」
   - 「渠道变革」→ 「直播间美妆达人打光下的护肤品陈列，暖光氛围」
3. 必须包含的画质词：真实感，高端商业产品级摄影，浅景深，精准布光，8K 分辨率，电影级色调。
4. 绝对禁止画面出现：文字、字母、数字、图表坐标轴、按钮、任何 UI 元素。
5. 画面必须传递出「高端、洁净、专业、值得信赖」的品牌质感。"""

AGENT3_USER = """请将以下场景描述转化为豆包星绘能理解的高品质摄影提示词。

画面意图：{image_intent}
所属章节语境：{chapter_title} — {core_proposition}

请直接输出用于 API 调用的纯中文逗号分隔提示词，不超过 120 字，不含任何标点之外的符号："""

# ========================
# 角色四：图表数据官
# ========================
AGENT4_SYSTEM = """# Role: 数据可视化官

你是一位同时精通数据分析与 ECharts 图表配置的专家。
你只做一件事：从已完成的白皮书章节正文中，
精准提炼出最具说服力的数据关系，并将其转化为
可直接被 ECharts 渲染的标准图表配置 JSON。

# 绝对红线:
1. 只能输出合法 JSON，不含 ```json 标记，不含任何废话。
2. type 只能从以下选择：bar / line / pie / area / scatter
3. labels 数量：3 ~ 7 个，不得过少或过多。
4. 如有多个数据系列，datasets 数组长度不超过 3 个，防止图表信息过载。
5. 所有数值必须量级合理，与正文语境一致（例如正文说「破千亿」，数值就不能是个位数）。
6. title 字段：12 字以内，直接点明图表核心结论，而非只描述数据维度。"""

AGENT4_USER = """白皮书标题：《{title}》
当前章节：【{chapter_title}】
图表意图：{chart_intent}

章节正文：
<CHAPTER_CONTENT>
{chapter_markdown}
</CHAPTER_CONTENT>

根据上方正文内容，生成一份 ECharts 图表数据配置。
数据必须来源于或高度吻合正文中的论述，不得凭空捏造与正文矛盾的数字。

输出 JSON Schema：
{{
  "type": "bar",
  "title": "图表核心结论标题，12字以内",
  "labels": ["标签1", "标签2", "标签3"],
  "datasets": [
    {{
      "name": "系列名称",
      "values": [100, 200, 150]
    }}
  ]
}}

请立即输出 JSON："""


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def _parse_llm_json(raw: str) -> dict:
    """从 LLM 输出中健壮地提取 JSON，自动修复常见问题"""
    clean = raw.strip()
    # 去掉 markdown 代码块
    if clean.startswith("```"):
        clean = clean.split("\n", 1)[-1]
    if clean.endswith("```"):
        clean = clean.rsplit("```", 1)[0]
    clean = clean.strip()

    # 第一次尝试：直接解析
    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        pass

    # 第二次尝试：提取最外层 { ... }
    start = clean.find("{")
    end = clean.rfind("}") + 1
    if start >= 0 and end > start:
        fragment = clean[start:end]
        try:
            return json.loads(fragment)
        except json.JSONDecodeError:
            pass

        # 第三次尝试：修复常见 LLM JSON 错误
        fixed = fragment
        # 修复尾部多余逗号 (trailing comma before } or ])
        fixed = re.sub(r',\s*([}\]])', r'\1', fixed)
        # 修复中文引号
        fixed = fixed.replace('\u201c', '"').replace('\u201d', '"')
        fixed = fixed.replace('\u2018', "'").replace('\u2019', "'")
        # 修复未转义的换行符在字符串值内
        fixed = re.sub(r'(?<=": ")(.*?)(?=")', lambda m: m.group(0).replace('\n', '\\n'), fixed)
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            pass

    # 第四次尝试：修复被截断的 JSON（补全缺失的括号）
    candidate = clean[start:] if start >= 0 else clean
    # 扫描：跟踪字符串状态和括号栈
    in_str = False
    last_safe = 0  # 最后一个结构完整的位置
    stack = []     # 记录 { [ 的嵌套顺序
    i = 0
    while i < len(candidate):
        ch = candidate[i]
        if ch == '\\' and in_str:
            i += 2
            continue
        if ch == '"':
            in_str = not in_str
            if not in_str:
                last_safe = i + 1
        elif not in_str:
            if ch in ('{', '['):
                stack.append('}' if ch == '{' else ']')
                last_safe = i + 1
            elif ch in ('}', ']'):
                if stack:
                    stack.pop()
                last_safe = i + 1
            else:
                # 数字、布尔值、逗号、冒号等也是安全位置
                last_safe = i + 1
        i += 1
    if in_str or stack:
        # 回退到最后安全位置，然后补全
        candidate = candidate[:last_safe]
        candidate = re.sub(r',\s*$', '', candidate)
        candidate = re.sub(r',\s*"[^"]*"\s*:\s*$', '', candidate)
        # 重新计算剩余未闭合括号
        stack2 = []
        in_str2 = False
        for ch in candidate:
            if ch == '"' and (not candidate or True):
                in_str2 = not in_str2
            elif not in_str2:
                if ch in ('{', '['):
                    stack2.append('}' if ch == '{' else ']')
                elif ch in ('}', ']') and stack2:
                    stack2.pop()
        # 按嵌套逆序补全
        candidate += ''.join(reversed(stack2))
        candidate = re.sub(r',\s*([}\]])', r'\1', candidate)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            pass

    # 全部失败，抛出有用的错误信息
    preview = clean[:200] + '...' if len(clean) > 200 else clean
    raise ValueError(f"无法解析 LLM 返回的 JSON。前200字符: {preview}")


def _record_pipeline_failure(stage: str, error: Exception, raw_preview: str = "") -> None:
    try:
        os.makedirs(os.path.dirname(PIPELINE_MEMORY_PATH) or ".", exist_ok=True)
        records = []
        if os.path.exists(PIPELINE_MEMORY_PATH):
            with open(PIPELINE_MEMORY_PATH, "r", encoding="utf-8") as f:
                loaded = json.load(f)
                if isinstance(loaded, list):
                    records = loaded
        error_text = str(error)
        lesson = "输出必须是完整合法 JSON，不要追加解释文字，所有对象和数组必须闭合。"
        if "无法解析" not in error_text:
            lesson = "上一轮在该阶段失败，下一轮需要优先规避同类错误。"
        records.append(
            {
                "timestamp": int(time.time()),
                "stage": stage,
                "error": error_text[:500],
                "raw_preview": raw_preview[:500],
                "lesson": lesson,
            }
        )
        with open(PIPELINE_MEMORY_PATH, "w", encoding="utf-8") as f:
            json.dump(records[-30:], f, ensure_ascii=False, indent=2)
    except Exception as memory_error:
        print(f"[Memory] 记录失败原因失败: {memory_error}")


def _load_pipeline_lessons(limit: int = 3) -> str:
    try:
        if not os.path.exists(PIPELINE_MEMORY_PATH):
            return ""
        with open(PIPELINE_MEMORY_PATH, "r", encoding="utf-8") as f:
            records = json.load(f)
        if not isinstance(records, list):
            return ""
        recent = records[-limit:]
        if not recent:
            return ""
        lines = ["以下是系统上次失败记忆，请在本次生成中主动规避："]
        for item in recent:
            lines.append(
                f"- 阶段: {item.get('stage', 'unknown')}；失败: {item.get('error', '')[:120]}；修正: {item.get('lesson', '')}"
            )
        return "\n".join(lines)
    except Exception as memory_error:
        print(f"[Memory] 读取失败记忆失败: {memory_error}")
        return ""


# ---------------------------------------------------------------------------
# 四阶段生成管线
# ---------------------------------------------------------------------------
def _append_reference_links(markdown: str, references: List[Dict[str, str]]) -> str:
    if not references:
        return markdown

    ref_lines = ["## 参考资料与链接\n"]
    for idx, ref in enumerate(references, 1):
        title = ref.get("title", "未命名来源")
        url = ref.get("url", "")
        publisher = ref.get("publisher", "未知发布方")
        source_type = ref.get("source_type", "media")
        published_at = ref.get("published_at") or "日期未提取"
        ref_lines.append(
            f"{idx}. [{title}]({url})  \n"
            f"   发布方：{publisher} ｜ 类型：{source_type} ｜ 日期：{published_at}"
        )

    return markdown + "\n\n---\n\n" + "\n".join(ref_lines)


def _clean_chart_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _sanitize_placeholder_numbers(markdown: str) -> str:
    """清理模型在资料不足时吐出的 X/XX/XXX 占位数字。"""
    if not markdown:
        return markdown
    text = markdown
    replacements = [
        (r"X{1,4}\s*亿元", "公开资料未披露具体数值"),
        (r"X{1,4}\s*万元", "公开资料未披露具体数值"),
        (r"X{1,4}\s*元", "公开资料未披露具体数值"),
        (r"X{1,4}\s*%", "公开资料未披露具体比例"),
        (r"XX+\.?X*", "公开资料未披露具体数值"),
    ]
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    text = re.sub(r"(突破|达到|约|超过)\s*公开资料未披露具体数值", "公开资料未披露具体数值", text)
    return text


def _image_prompt_cache_key(image_intent: str, chapter_title: str, core_proposition: str) -> str:
    raw = "\n".join([image_intent or "", chapter_title or "", core_proposition or ""])
    normalized = re.sub(r"\s+", " ", raw.strip().lower())
    return hashlib.sha1(normalized.encode("utf-8")).hexdigest()


def _load_image_prompt_cache(cache_key: str) -> Optional[str]:
    try:
        if IMAGE_PROMPT_CACHE_TTL <= 0 or not os.path.exists(IMAGE_PROMPT_CACHE_PATH):
            return None
        with open(IMAGE_PROMPT_CACHE_PATH, "r", encoding="utf-8") as f:
            records = json.load(f)
        if not isinstance(records, dict):
            return None
        item = records.get(cache_key)
        if not isinstance(item, dict):
            return None
        if int(time.time()) - int(item.get("timestamp", 0) or 0) > IMAGE_PROMPT_CACHE_TTL:
            return None
        prompt = item.get("prompt")
        return prompt if isinstance(prompt, str) and prompt.strip() else None
    except Exception as exc:
        print(f"[Pipeline][ImagePromptCache] 读取失败: {exc}")
        return None


def _save_image_prompt_cache(cache_key: str, prompt: str) -> None:
    try:
        if IMAGE_PROMPT_CACHE_TTL <= 0 or not prompt:
            return
        os.makedirs(os.path.dirname(IMAGE_PROMPT_CACHE_PATH) or ".", exist_ok=True)
        records: Dict[str, Any] = {}
        if os.path.exists(IMAGE_PROMPT_CACHE_PATH):
            with open(IMAGE_PROMPT_CACHE_PATH, "r", encoding="utf-8") as f:
                loaded = json.load(f)
                if isinstance(loaded, dict):
                    records = loaded
        records[cache_key] = {"timestamp": int(time.time()), "prompt": prompt}
        records = dict(list(records.items())[-200:])
        with open(IMAGE_PROMPT_CACHE_PATH, "w", encoding="utf-8") as f:
            json.dump(records, f, ensure_ascii=False, indent=2)
    except Exception as exc:
        print(f"[Pipeline][ImagePromptCache] 写入失败: {exc}")


def _image_intent_enabled(image_intent: str) -> bool:
    value = str(image_intent or "").strip().lower()
    return bool(value and value not in {"none", "无", "无需", "不需要", "null"})


def _build_fallback_image_intent(chapter: Dict[str, Any], title: str) -> str:
    case_direction = str(chapter.get("case_direction") or "").strip()
    chapter_title = str(chapter.get("chapter_title") or "").strip()
    core = str(chapter.get("core_proposition") or "").strip()
    topic = case_direction or chapter_title or core or title
    return (
        f"围绕「{topic}」拍摄真实美妆行业商业场景："
        "高端产品陈列、门店柜台或实验室质检台面，体现品牌、品类、消费者洞察与专业质感"
    )


def _ensure_minimum_image_intents(
    chapters: List[Dict[str, Any]],
    title: str,
    min_count: int = 3,
    max_count: int = 4,
) -> None:
    """大纲过度保守时补足配图意图，避免开启配图能力后整篇无图。"""
    current = [
        ch for ch in chapters
        if _image_intent_enabled(ch.get("image_intent", "none"))
    ]
    if len(current) >= min_count:
        return

    needed = min(max_count, max(min_count - len(current), 0))
    if needed <= 0:
        return

    def chapter_priority(item: Tuple[int, Dict[str, Any]]) -> Tuple[int, int]:
        idx, ch = item
        text = " ".join(
            str(ch.get(key) or "")
            for key in ("visual_type", "case_direction", "chapter_title", "core_proposition", "content_guidelines")
        )
        case_score = 2 if any(marker in text for marker in ("案例", "场景", "门店", "渠道", "产品", "消费者", "实验室", "柜台")) else 0
        non_chart_score = 1 if str(ch.get("visual_type") or "").lower() != "chart" else 0
        return (case_score + non_chart_score, -idx)

    candidates = [
        item for item in enumerate(chapters)
        if not _image_intent_enabled(item[1].get("image_intent", "none"))
    ]
    candidates.sort(key=chapter_priority, reverse=True)
    for _, chapter in candidates[:needed]:
        chapter["image_intent"] = _build_fallback_image_intent(chapter, title)
        if str(chapter.get("visual_type") or "").strip().lower() in {"", "none", "chart"}:
            chapter["visual_type"] = "case_photo"
        _progress_print(
            "[Pipeline][ImageFallback] "
            f"chapter={chapter.get('chapter_id')} image_intent={chapter.get('image_intent')[:80]}",
            stage="pipeline",
        )


def _normalize_chart_config(raw_chart: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """把 LLM 图表 JSON 收敛成前端 ECharts 可安全渲染的最小结构。"""
    if not isinstance(raw_chart, dict):
        return None

    chart_type = str(raw_chart.get("type") or "bar").strip().lower()
    if chart_type not in {"bar", "line", "pie", "area", "scatter"}:
        chart_type = "bar"

    labels = raw_chart.get("labels")
    datasets = raw_chart.get("datasets")
    if not isinstance(labels, list) or not isinstance(datasets, list) or not datasets:
        return None

    safe_labels = [_clean_chart_text(label)[:24] for label in labels if str(label).strip()]
    safe_labels = safe_labels[:7]
    if len(safe_labels) < 2:
        return None

    safe_datasets = []
    for idx, dataset in enumerate(datasets[:3]):
        if not isinstance(dataset, dict):
            continue
        values = dataset.get("values")
        if not isinstance(values, list):
            continue
        numeric_values = []
        for value in values[:len(safe_labels)]:
            try:
                numeric_values.append(float(value))
            except (TypeError, ValueError):
                numeric_values.append(0.0)
        if len(numeric_values) < len(safe_labels):
            numeric_values.extend([0.0] * (len(safe_labels) - len(numeric_values)))
        safe_datasets.append(
            {
                "name": _clean_chart_text(dataset.get("name") or f"系列{idx + 1}")[:24],
                "values": numeric_values,
            }
        )
    if not safe_datasets:
        return None

    return {
        "type": chart_type,
        "title": _clean_chart_text(raw_chart.get("title") or "数据趋势")[:24],
        "labels": safe_labels,
        "datasets": safe_datasets,
    }


def run_four_agent_pipeline(
    data_text: str,
    research_context: str = "",
    references: Optional[List[Dict[str, str]]] = None,
) -> str:
    """
    Agent1(战略官) → [Agent2(主笔) + Agent3(视觉导演)] 并发
    → Agent4(图表数据官) 并发 → 合并完整白皮书 Markdown
    新版：支持 front_matter / back_matter / sub_sections 等完整白皮书体例
    """
    llm = get_llm_provider()

    # ===== 第一批：战略官 生成大纲（串行，带重试）=====
    _progress_print("[Pipeline] 第一批：战略官生成大纲...", stage="pipeline")
    outline = None
    for attempt in range(3):
        outline_raw = ""
        try:
            outline_prompt = AGENT1_USER.format(data=data_text)
            failure_lessons = _load_pipeline_lessons()
            if failure_lessons:
                outline_prompt += f"\n\n<PIPELINE_FAILURE_MEMORY>\n{failure_lessons}\n</PIPELINE_FAILURE_MEMORY>"
            if research_context:
                outline_prompt += f"\n\n<RECENT_WEB_RESEARCH>\n{research_context}\n</RECENT_WEB_RESEARCH>"
            outline_raw = llm.generate(
                prompt=outline_prompt,
                system=AGENT1_SYSTEM,
            )
            outline = _parse_llm_json(outline_raw)
            break
        except Exception as e:
            _progress_print(f"[Pipeline] 战略官第 {attempt+1} 次尝试失败: {e}", stage="pipeline")
            _record_pipeline_failure("outline_json", e, outline_raw)
            if attempt == 2:
                raise ValueError(f"战略官连续 3 次生成大纲失败: {e}")
    meta = outline.get("whitepaper_meta", {})
    front_matter = outline.get("front_matter", {})
    back_matter = outline.get("back_matter", {})
    chapters = outline.get("chapters", [])
    if not chapters:
        raise ValueError("战略官未生成任何章节")

    title = meta.get("title", "CIBE 美业数据洞察白皮书")
    subtitle = meta.get("subtitle", "")
    tone = meta.get("global_tone", "专业、数据驱动")
    target_audience = meta.get("target_audience", "行业决策层、渠道商、投资机构")
    _progress_print(f"[Pipeline] 大纲完成：{title}，共 {len(chapters)} 章", stage="pipeline")
    if image_generator.available:
        _ensure_minimum_image_intents(chapters, title)

    # ===== 第二批：主笔 + 视觉导演（并发）=====
    _progress_print("[Pipeline] 第二批：主笔 + 视觉导演并发...", stage="pipeline")
    chapter_texts = {}   # chapter_id → markdown
    image_prompts = {}   # chapter_id → 摄影提示词

    def _write_chapter(ch, ch_idx):
        CN_NUMS = "一二三四五六七八九十"
        ch_number = CN_NUMS[ch_idx - 1] if ch_idx <= 10 else str(ch_idx)
        ch_id = ch["chapter_id"]
        chart_instr = (
            f'在最合适的段落后插入：[CHART_PLACEHOLDER id="{ch_id}"]'
            if ch.get("needs_chart") else "（本章无需图表）"
        )
        # image_intent 不为 none 就需要配图
        img_intent = ch.get("image_intent", "none")
        needs_img = _image_intent_enabled(img_intent)
        image_instr = (
            f'在最合适的案例/场景段落后插入：[IMAGE_PLACEHOLDER id="{ch_id}"]'
            if needs_img else "（本章无需配图）"
        )
        # 构造二级小节文本
        sub_sections = ch.get("sub_sections", [])
        sub_text = "\n".join(
            f"- {s.get('section_id','')}: {s.get('section_title','')} — {s.get('section_focus','')}"
            for s in sub_sections
        ) if sub_sections else "（由你自行组织二级结构）"
        # 构造数据锚点文本
        anchors = ch.get("data_anchors", [])
        anchors_text = "\n".join(f"- {a}" for a in anchors) if anchors else "（从正文内容中自行提取关键数据）"

        md = llm.generate(
            prompt=(
                AGENT2_USER.format(
                title=title,
                tone=tone,
                target_audience=target_audience,
                chapter_id=ch_id,
                chapter_title=ch.get("chapter_title", ""),
                core_proposition=ch.get("core_proposition", ""),
                chapter_intro=ch.get("chapter_intro", ""),
                transition=ch.get("transition_from_previous", ""),
                guidelines=ch.get("content_guidelines", ""),
                sub_sections_text=sub_text,
                data_anchors_text=anchors_text,
                case_direction=ch.get("case_direction", ""),
                chapter_summary=ch.get("chapter_summary", ""),
                chart_placeholder_instruction=chart_instr,
                image_placeholder_instruction=image_instr,
                chapter_number=ch_number,
                )
                + (
                    f"\n\n<RECENT_WEB_RESEARCH>\n{research_context}\n</RECENT_WEB_RESEARCH>"
                    if research_context else ""
                )
            ),
            system=AGENT2_SYSTEM,
        )
        return ("chapter", ch_id, md.strip())

    def _gen_image_prompt(ch):
        ch_id = ch["chapter_id"]
        cache_key = _image_prompt_cache_key(
            ch.get("image_intent", ""),
            ch.get("chapter_title", ""),
            ch.get("core_proposition", ""),
        )
        cached_prompt = _load_image_prompt_cache(cache_key)
        if cached_prompt:
            _progress_print(f"[Pipeline][ImagePromptCache] 命中 chapter={ch_id}", stage="pipeline")
            return ("image", ch_id, cached_prompt)
        prompt_text = llm.generate(
            prompt=AGENT3_USER.format(
                image_intent=ch.get("image_intent", ""),
                chapter_title=ch.get("chapter_title", ""),
                core_proposition=ch.get("core_proposition", ""),
            ),
            system=AGENT3_SYSTEM,
        )
        prompt_text = prompt_text.strip().strip('"').strip("'")
        _save_image_prompt_cache(cache_key, prompt_text)
        return ("image", ch_id, prompt_text)

    with ThreadPoolExecutor(max_workers=3) as pool:
        futures = []
        for ch_idx, ch in enumerate(chapters, 1):
            futures.append(pool.submit(_write_chapter, ch, ch_idx))
            # image_intent 不为 none 就生成配图（不受 visual_type 限制）
            img_intent = ch.get("image_intent", "none")
            if image_generator.available and _image_intent_enabled(img_intent):
                futures.append(pool.submit(_gen_image_prompt, ch))

        for fut in as_completed(futures):
            try:
                kind, ch_id, result = fut.result()
                if kind == "chapter":
                    chapter_texts[ch_id] = result
                elif kind == "image":
                    image_prompts[ch_id] = result
            except Exception as e:
                _progress_print(f"[Pipeline] 第二批任务失败: {e}", stage="pipeline")

    _progress_print(f"[Pipeline] 第二批完成：{len(chapter_texts)} 章正文，{len(image_prompts)} 个配图词", stage="pipeline")

    # ===== 第三批：图表数据官（并发）=====
    _progress_print("[Pipeline] 第三批：图表数据官并发...", stage="pipeline")
    chart_data = {}  # chapter_id → chart dict

    def _gen_chart(ch):
        ch_id = ch["chapter_id"]
        ch_md = chapter_texts.get(ch_id, "")
        if not ch_md:
            return ("chart", ch_id, None)
        raw = llm.generate(
            prompt=AGENT4_USER.format(
                title=title,
                chapter_title=ch.get("chapter_title", ""),
                chart_intent=ch.get("chart_intent", ""),
                chapter_markdown=ch_md,
            ),
            system=AGENT4_SYSTEM,
        )
        return ("chart", ch_id, _parse_llm_json(raw))

    with ThreadPoolExecutor(max_workers=3) as pool:
        chart_futures = []
        for ch in chapters:
            if ch.get("needs_chart"):
                chart_futures.append(pool.submit(_gen_chart, ch))

        for fut in as_completed(chart_futures):
            try:
                _, ch_id, cdata = fut.result()
                if cdata:
                    chart_data[ch_id] = cdata
            except Exception as e:
                _progress_print(f"[Pipeline] 图表数据官失败: {e}", stage="pipeline")

    chart_data = {
        ch_id: normalized
        for ch_id, normalized in (
            (ch_id, _normalize_chart_config(cdata))
            for ch_id, cdata in chart_data.items()
        )
        if normalized
    }
    _progress_print(f"[Pipeline] 第三批完成：{len(chart_data)} 个图表", stage="pipeline")

    # ===== 合并阶段 =====
    _progress_print("[Pipeline] 合并完整白皮书...", stage="pipeline")

    # 清理正文中残留的括号建议/元注释
    _meta_comment_re = re.compile(
        r'[（(]\s*(?:图表建议|配图建议|建议|注|备注|说明|编辑批注|TODO)\s*[：:].+?[）)]\s*',
        re.DOTALL,
    )
    _star_comment_re = re.compile(
        r'\*\s*[（(]\s*(?:图表建议|配图建议|建议|注).+?[）)]\s*\*\s*',
        re.DOTALL,
    )

    # --- 封面 ---
    parts = []


    # --- 执行摘要 ---
    exec_summary = front_matter.get("executive_summary", {})
    if exec_summary:
        es_lines = ["## 执行摘要\n"]
        # 总判断
        judgement = exec_summary.get("executive_judgement", "")
        if judgement:
            es_lines.append(f"**核心判断：** {judgement}\n")
        # 关键发现
        findings = exec_summary.get("key_findings", [])
        if findings:
            es_lines.append("### 关键发现\n")
            for i, f in enumerate(findings, 1):
                es_lines.append(f"**{i}.** {f}\n")
        # 核心数据结论
        data_pts = exec_summary.get("core_data_points", [])
        if data_pts:
            es_lines.append("### 核心数据\n")
            for d in data_pts:
                es_lines.append(f"- {d}")
        parts.append("\n".join(es_lines))

    # --- 目录 ---
    toc_lines = ["## 目录\n"]
    toc_lines.append("- 执行摘要")
    toc_lines.append("- 研究说明")
    for i, ch in enumerate(chapters, 1):
        ch_title = ch.get("chapter_title", f"第{i}章")
        toc_lines.append(f"{i}. {ch_title}")
        for s in ch.get("sub_sections", []):
            toc_lines.append(f"   - {s.get('section_title', '')}")
    toc_lines.append("- 结论与行动建议")
    toc_lines.append("- 参考说明")
    parts.append("\n".join(toc_lines))

    # --- 研究说明 ---
    research = front_matter.get("research_note", {})
    if research:
        rn_lines = ["## 研究说明\n"]
        if research.get("research_object"):
            rn_lines.append(f"**研究对象：** {research['research_object']}\n")
        if research.get("time_scope"):
            rn_lines.append(f"**时间范围：** {research['time_scope']}\n")
        sources = research.get("data_sources", [])
        if sources:
            rn_lines.append(f"**数据来源：** {'、'.join(sources)}\n")
        if research.get("sample_scope"):
            rn_lines.append(f"**样本口径：** {research['sample_scope']}\n")
        if research.get("method_boundary"):
            rn_lines.append(f"**研究边界：** {research['method_boundary']}\n")
        if research.get("term_definition"):
            rn_lines.append(f"**名词定义：** {research['term_definition']}\n")
        parts.append("\n".join(rn_lines))

    # --- 正文章节 ---
    for ch in chapters:
        ch_id = ch.get("chapter_id", "")
        md = chapter_texts.get(ch_id, "")
        # 清理元注释
        md = _meta_comment_re.sub('', md)
        md = _star_comment_re.sub('', md)

        # 替换 [CHART_PLACEHOLDER id="XX"]
        if ch_id in chart_data:
            cd = chart_data[ch_id]
            chart_type = cd["type"]
            chart_title = html.escape(cd["title"], quote=True)
            data_obj = {"labels": cd["labels"], "datasets": cd["datasets"]}
            chart_payload = html.escape(json.dumps(data_obj, ensure_ascii=False), quote=True)
            chart_tag = f'<chart type="{chart_type}" title="{chart_title}" data="{chart_payload}"></chart>'
            md = re.sub(
                r'\[CHART_PLACEHOLDER\s+id="?' + re.escape(ch_id) + r'"?\]',
                chart_tag,
                md,
            )

        # 替换 [IMAGE_PLACEHOLDER id="XX"]
        if ch_id in image_prompts:
            image_prompt = html.escape(image_prompts[ch_id], quote=True)
            image_tag = f'<image prompt="{image_prompt}"></image>'
            md = re.sub(
                r'\[IMAGE_PLACEHOLDER\s+id="?' + re.escape(ch_id) + r'"?\]',
                image_tag,
                md,
            )

        parts.append(md)

    # --- 结论与行动建议 ---
    if back_matter:
        conclusion = back_matter.get("conclusion_summary", {})
        actions = back_matter.get("action_recommendations", {})
        risks = back_matter.get("risk_and_boundary_notes", [])

        bm_lines = ["## 结论与行动建议\n"]
        if conclusion.get("final_judgement"):
            bm_lines.append(f"**总结性判断：** {conclusion['final_judgement']}\n")
        implications = conclusion.get("industry_implications", [])
        if implications:
            bm_lines.append("### 行业启示\n")
            for imp in implications:
                bm_lines.append(f"- {imp}")
            bm_lines.append("")

        if actions:
            if actions.get("for_brands"):
                bm_lines.append("### 对品牌方的建议\n")
                for a in actions["for_brands"]:
                    bm_lines.append(f"- {a}")
                bm_lines.append("")
            if actions.get("for_channels"):
                bm_lines.append("### 对渠道方的建议\n")
                for a in actions["for_channels"]:
                    bm_lines.append(f"- {a}")
                bm_lines.append("")
            if actions.get("for_investment_or_strategy"):
                bm_lines.append("### 对投资/战略团队的建议\n")
                for a in actions["for_investment_or_strategy"]:
                    bm_lines.append(f"- {a}")
                bm_lines.append("")

        if risks:
            bm_lines.append("### 风险与边界提示\n")
            for r in risks:
                bm_lines.append(f"- {r}")
            bm_lines.append("")

        parts.append("\n".join(bm_lines))

    # --- 参考说明 ---
    refs = back_matter.get("references_or_appendix", {})
    if refs and refs.get("required"):
        ref_lines = ["## 参考说明与附录\n"]
        if refs.get("content_direction"):
            ref_lines.append(f"{refs['content_direction']}\n")
        parts.append("\n".join(ref_lines))

    # --- 拼合 ---
    full_markdown = "\n\n---\n\n".join(parts)
    full_markdown = _append_reference_links(full_markdown, references or [])
    _progress_print(f"[Pipeline] 白皮书生成完毕，总长度 {len(full_markdown)} 字符", stage="pipeline")
    
    # 如果启用反思功能，对生成的白皮书进行反思和验证
    if ENABLE_REFLECTION and _should_skip_reflection_fast_path(full_markdown, references):
        _progress_print(
            "[Pipeline] 跳过反思改写："
            f"chars={len(full_markdown)} refs={len(references or [])}",
            stage="pipeline",
        )
    elif ENABLE_REFLECTION:
        _progress_print("[Pipeline] 开始反思验证阶段...", stage="pipeline")
        # 注意：这里需要同步版本的反思验证
        full_markdown = _apply_reflection_and_verification_sync(full_markdown, data_text)
    
    return _sanitize_placeholder_numbers(full_markdown)


def _reflection_pass_threshold(reflection_mode: str) -> float:
    """System-generated reflections are scored against a lower calibration bar."""
    if reflection_mode in {"auto_generated", "source_consistency"}:
        return 0.35
    return 0.7


def _should_skip_reflection_fast_path(
    whitepaper_markdown: str,
    references: Optional[List[Dict[str, str]]] = None,
) -> bool:
    text = whitepaper_markdown or ""
    refs = references or []
    if len(text) < FAST_REFLECTION_MIN_CHARS:
        return False
    if len(refs) < FAST_REFLECTION_MIN_REFS:
        return False
    placeholder_markers = ("XXX", "XX%", "X亿元", "公开资料未披露具体数值")
    if any(marker in text for marker in placeholder_markers):
        return False
    if text.count("## ") < 4:
        return False
    return True


def _should_run_legacy_data_verification(reflection_mode: str) -> bool:
    """Avoid noisy number-by-number checks when source consistency already ran."""
    return reflection_mode != "source_consistency"


def _apply_reflection_and_verification_sync(whitepaper_markdown: str, original_data: str) -> str:
    """
    同步版本的反思和验证功能
    
    Args:
        whitepaper_markdown: 生成的白皮书内容
        original_data: 原始数据文本
        
    Returns:
        经过反思和验证后的白皮书内容
    """
    try:
        # 初始化反思代理
        reflection_agent = ReflectionAgent(
            storage_path=REFLECTION_STORAGE_PATH,
            criteria_weights={
                "clarity_coherence": 0.25,
                "depth_analysis": 0.25,
                "actionability": 0.20,
                "self_awareness": 0.20,
                "structure_organization": 0.10
            }
        )
        
        # 创建反思会话
        session_id = f"whitepaper_{int(time.time())}"
        reflection_session = ReflectionSession(
            session_id=session_id,
            query="白皮书生成质量评估",
            response=whitepaper_markdown,
            reflection="",
            evaluation_result=None
        )
        
        # 执行反思评估
        print(f"[Reflection] 开始反思评估会话 {session_id}")
        result_session = reflection_agent.evaluate_and_reflect(
            reflection_session.query,
            reflection_session.response,
            reflection_session.reflection,
            session_id=session_id
        )
        evaluation_result = result_session.evaluation_result
        reflection_mode = result_session.reflection_mode
        pass_threshold = _reflection_pass_threshold(reflection_mode)
        
        # 评估反思质量
        if evaluation_result.overall_score >= pass_threshold:
            print(f"[Reflection] 反思质量良好，评分: {evaluation_result.overall_score:.2f}")
            # 如果质量良好，添加反思总结到白皮书
            reflection_summary = f"""
## 反思与验证

本白皮书经过了AI反思系统的质量评估，评分: {evaluation_result.overall_score:.2f}/1.00

### 主要优势
{chr(10).join(f"- {strength}" for strength in evaluation_result.strengths)}

### 改进建议
{chr(10).join(f"- {recommendation}" for recommendation in evaluation_result.recommendations)}

### 验证过程
本报告通过AI反思系统进行了多维度评估，包括清晰度、分析深度、行动导向性、自我意识和结构组织等方面。

---
"""
            whitepaper_markdown += reflection_summary
            
        else:
            print(f"[Reflection] 反思质量不足，评分: {evaluation_result.overall_score:.2f}，需要重新生成")
            
            # 如果质量不足，记录原因
            reflection_session.reflection = f"""
反思评估结果:
- 总体评分: {evaluation_result.overall_score:.2f}
- 主要问题: {', '.join(evaluation_result.weaknesses)}
- 改进建议: {', '.join(evaluation_result.recommendations)}
"""
            
            reflection_session.evaluation_result = evaluation_result
            # evaluate_and_reflect 内部已自动保存，无需再次调用

            # 生成改进后的版本
            print("[Reflection] 生成改进版本...")
            improved_markdown = _generate_improved_whitepaper_sync(
                whitepaper_markdown, 
                original_data, 
                evaluation_result
            )
            
            # 再次验证改进版本
            improved_result_session = reflection_agent.evaluate_and_reflect(
                reflection_session.query,
                improved_markdown,
                f"基于前次评估结果的改进版本。原问题: {', '.join(evaluation_result.weaknesses)}",
                session_id=f"{session_id}_improved"
            )
            improved_evaluation = improved_result_session.evaluation_result
            
            reflection_mode = improved_result_session.reflection_mode
            improved_threshold = _reflection_pass_threshold(reflection_mode)
            if improved_evaluation.overall_score >= improved_threshold:
                print(f"[Reflection] 改进版本通过验证，评分: {improved_evaluation.overall_score:.2f}")
                
                # 添加改进说明
                improvement_summary = f"""
## 反思与验证

本白皮书经过AI反思系统的多轮质量评估和改进，最终评分: {improved_evaluation.overall_score:.2f}/1.00

### 改进过程
1. **初始版本评分**: {evaluation_result.overall_score:.2f}
2. **识别问题**: {', '.join(evaluation_result.weaknesses)}
3. **改进措施**: {', '.join(evaluation_result.recommendations)}
4. **最终版本评分**: {improved_evaluation.overall_score:.2f}

### 最终评估结果
- **主要优势**: {', '.join(improved_evaluation.strengths)}
- **持续改进**: {', '.join(improved_evaluation.recommendations)}

---
"""
                improved_markdown += improvement_summary
                whitepaper_markdown = improved_markdown
            else:
                print(f"[Reflection] 改进版本仍未达到要求，评分: {improved_evaluation.overall_score:.2f}")
                # 保存最终评估结果
                # evaluate_and_reflect 内部已自动保存改进后评估，无需再次调用
        
        # 如果启用网络搜索，对关键数据进行验证
        if (
            ENABLE_WEB_SEARCH
            and ENABLE_DATA_VERIFICATION
            and _should_run_legacy_data_verification(reflection_mode)
        ):
            print("[Verification] 开始数据验证...")
            whitepaper_markdown = _verify_data_with_web_search_sync(whitepaper_markdown)
        
        return whitepaper_markdown
        
    except Exception as e:
        print(f"[Reflection] 反思验证过程中出现错误: {e}")
        # 如果反思验证失败，返回原始白皮书
        return whitepaper_markdown


def _generate_improved_whitepaper_sync(original_markdown: str, original_data: str, evaluation_result) -> str:
    """
    同步版本的改进白皮书生成
    
    Args:
        original_markdown: 原始白皮书内容
        original_data: 原始数据
        evaluation_result: 反思评估结果
        
    Returns:
        改进后的白皮书内容
    """
    try:
        llm = get_llm_provider()
        
        # 构建改进提示
        improvement_prompt = f"""
基于以下反思评估结果，请对白皮书进行改进：

**原始白皮书**:
{original_markdown}

**反思评估结果**:
- 总体评分: {evaluation_result.overall_score:.2f}
- 主要优势: {', '.join(evaluation_result.strengths)}
- 主要问题: {', '.join(evaluation_result.weaknesses)}
- 改进建议: {', '.join(evaluation_result.recommendations)}

**改进要求**:
1. 针对评估中发现的问题进行具体改进
2. 保持原有的结构和主要内容
3. 提升分析的深度和逻辑性
4. 增强内容的可读性和实用性
5. 确保数据准确性和论证的严密性

请生成改进后的白皮书，重点关注以下方面：
{chr(10).join(f"- {weakness}" for weakness in evaluation_result.weaknesses)}
"""
        
        # 使用LLM生成改进版本
        improved_markdown = llm.generate(
            prompt=improvement_prompt,
            system="""你是一位专业的白皮书编辑和质量改进专家。
你的任务是根据反思评估结果对白皮书进行改进。
请保持原有的核心内容和结构，但针对评估中发现的问题进行优化。
重点关注：
- 提升逻辑性和连贯性
- 增强分析的深度
- 改进语言表达
- 确保数据准确性
- 增强实用性建议

请直接输出改进后的白皮书内容，不要包含额外的说明。"""
        )
        
        return improved_markdown
        
    except Exception as e:
        print(f"[Improvement] 生成改进版本失败: {e}")
        return original_markdown


def _verify_data_with_web_search_sync(whitepaper_markdown: str) -> str:
    """
    同步版本的网络数据验证
    
    Args:
        whitepaper_markdown: 白皮书内容
        
    Returns:
        包含数据验证结果的白皮书
    """
    try:
        web_search = WebSearch(timeout=WEB_SEARCH_TIMEOUT)
        
        # 提取可能需要验证的数据点：必须带单位/语境，避免把孤零零的整数、年份
        # 当作"数据点"批量丢进验证（那只会刷一堆"需要人工审核"的噪声）。
        data_patterns = [
            r'\d+(?:\.\d+)?\s*(?:亿元|亿|万元|万|千万|百万|million|billion)',
            r'\d+(?:\.\d+)?\s*%',
            r'(?:同比|环比|增长|下降|上升|复合增长率|CAGR)\s*\d+(?:\.\d+)?\s*%?',
        ]
        
        verification_results = []
        
        for pattern in data_patterns:
            matches = re.finditer(pattern, whitepaper_markdown)
            for match in matches:
                data_point = match.group()
                context = whitepaper_markdown[max(0, match.start()-100):match.start()+100]
                
                # 验证数据点
                verification = web_search.verify_data_point(data_point, context)
                
                if verification["verified"] and verification["confidence"] >= 60:
                    verification_results.append({
                        "data_point": data_point,
                        "status": "verified",
                        "confidence": verification["confidence"],
                        "sources": len(verification["supporting_sources"])
                    })
                    print(f"[Verification] 数据点 '{data_point}' 已验证，置信度: {verification['confidence']}")
                else:
                    verification_results.append({
                        "data_point": data_point,
                        "status": "needs_review",
                        "confidence": verification["confidence"],
                        "message": verification.get("message", "验证结果不确定")
                    })
                    print(f"[Verification] 数据点 '{data_point}' 需要人工审核，置信度: {verification['confidence']}")
        
        # 验证结果只记录到后端日志，供开发者核对；不再写进白皮书正文。
        if verification_results:
            verified_cnt = sum(1 for r in verification_results if r["status"] == "verified")
            review_cnt = sum(1 for r in verification_results if r["status"] == "needs_review")
            print(
                f"[Verification] 汇总：共 {len(verification_results)} 个数据点，"
                f"已验证 {verified_cnt}，需人工核对 {review_cnt}（仅后端记录，不写入报告）"
            )

        return whitepaper_markdown
        
    except Exception as e:
        print(f"[Verification] 数据验证失败: {e}")
        return whitepaper_markdown


# ---------------------------------------------------------------------------
# API 路由
# ---------------------------------------------------------------------------

@app.get("/")
async def serve_index():
    """提供前端页面"""
    return FileResponse("index.html")


@app.get("/api/health")
async def health_check():
    return {
        "status": "ok",
        "deepseek": bool(DEEPSEEK_API_KEY),
        "image_gen": image_generator.available,
        "reflection_enabled": ENABLE_REFLECTION,
        "web_search_enabled": ENABLE_WEB_SEARCH,
        "data_verification_enabled": ENABLE_DATA_VERIFICATION,
        "timestamp": int(time.time()),
    }


@app.get("/api/auth/me")
async def auth_me(request: Request):
    username = _verify_session_token(request.cookies.get(AUTH_COOKIE_NAME, "")) if AUTH_ENABLED else AUTH_USERNAME
    user = _get_user(username) if username else None
    return {
        "authenticated": bool(username),
        "username": username or "",
        "role": (user or {}).get("role", ""),
        "auth_enabled": AUTH_ENABLED,
    }


@app.post("/api/auth/login")
async def auth_login(username: str = Form(...), password: str = Form(...)):
    if not AUTH_ENABLED:
        return {"status": "success", "username": AUTH_USERNAME}
    _init_app_db()
    valid_db_user = _verify_user_password(username, password)
    valid_env_user = (
        hmac.compare_digest(username, AUTH_USERNAME)
        and hmac.compare_digest(password, AUTH_PASSWORD)
    )
    if not (valid_db_user or valid_env_user):
        raise HTTPException(status_code=401, detail="账号或密码错误")
    response = JSONResponse({"status": "success", "username": username})
    response.set_cookie(
        AUTH_COOKIE_NAME,
        _create_session_token(username),
        max_age=AUTH_SESSION_TTL,
        httponly=True,
        samesite="lax",
        secure=os.getenv("AUTH_COOKIE_SECURE", "false").lower() == "true",
    )
    return response


@app.post("/api/auth/logout")
async def auth_logout():
    response = JSONResponse({"status": "success"})
    response.delete_cookie(AUTH_COOKIE_NAME)
    return response


@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    return Response(status_code=204)


@app.get("/api/progress/{job_id}")
async def progress_stream(job_id: str):
    async def event_generator():
        sent = 0
        idle_ticks = 0
        while True:
            with _PROGRESS_LOCK:
                events = list(_PROGRESS_EVENTS.get(job_id, []))
                done = _PROGRESS_DONE.get(job_id, False)
            for event in events[sent:]:
                sent += 1
                payload = json.dumps(event, ensure_ascii=False)
                event_name = "done" if event.get("stage") == "done" else "progress"
                yield f"event: {event_name}\ndata: {payload}\n\n"
            if done and sent >= len(events):
                break
            idle_ticks += 1
            if idle_ticks % 20 == 0:
                yield ": keep-alive\n\n"
            await asyncio.sleep(0.35)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


def _current_username(request: Request) -> str:
    if not AUTH_ENABLED:
        return AUTH_USERNAME
    return _verify_session_token(request.cookies.get(AUTH_COOKIE_NAME, "")) or AUTH_USERNAME


def _current_user(request: Request) -> Dict[str, Any]:
    username = _current_username(request)
    return _get_user(username) or {
        "username": username,
        "role": "admin" if username == AUTH_USERNAME else "user",
        "created_at": None,
    }


def _require_admin(request: Request) -> Dict[str, Any]:
    user = _current_user(request)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可执行此操作")
    return user


@app.get("/api/jobs")
async def list_generation_jobs(request: Request, limit: int = Query(20, ge=1, le=100)):
    username = _current_username(request)
    jobs = _list_generation_jobs(username, limit=limit)
    return {
        "status": "success",
        "jobs": [
            {
                **job,
                "title": _extract_job_title(job),
            }
            for job in jobs
        ],
    }


@app.get("/api/jobs/{job_id}")
async def get_generation_job(job_id: str, request: Request):
    username = _current_username(request)
    job = _get_generation_job(job_id)
    if not job or job.get("username") != username:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {
        "status": "success",
        "job": {
            "id": job["id"],
            "title": _extract_job_title(job),
            "status": job["status"],
            "markdown": job.get("markdown"),
            "error": job.get("error"),
            "created_at": job.get("created_at"),
            "started_at": job.get("started_at"),
            "finished_at": job.get("finished_at"),
        },
    }


@app.delete("/api/jobs/{job_id}")
async def delete_generation_job(job_id: str, request: Request):
    username = _current_username(request)
    job = _get_generation_job(job_id)
    if not job or job.get("username") != username:
        raise HTTPException(status_code=404, detail="任务不存在")
    _delete_generation_job(job_id)
    with _PROGRESS_LOCK:
        _PROGRESS_EVENTS.pop(job_id, None)
        _PROGRESS_DONE.pop(job_id, None)
    return {"status": "success", "job_id": job_id}


@app.get("/api/admin/users")
async def admin_list_users(request: Request, limit: int = Query(100, ge=1, le=200)):
    _require_admin(request)
    return {"status": "success", "users": _list_users(limit=limit)}


@app.get("/api/admin/dashboard")
async def admin_dashboard(request: Request, limit: int = Query(30, ge=1, le=100)):
    _require_admin(request)
    return {"status": "success", **_build_admin_dashboard(limit=limit)}


@app.get("/api/admin/jobs/{job_id}")
async def admin_get_job(job_id: str, request: Request):
    _require_admin(request)
    job = _get_generation_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {
        "status": "success",
        "job": {
            "id": job["id"],
            "username": job["username"],
            "status": job["status"],
            "markdown": job.get("markdown"),
            "data_text": job.get("data_text"),
            "error": job.get("error"),
            "created_at": job.get("created_at"),
            "started_at": job.get("started_at"),
            "finished_at": job.get("finished_at"),
            "title": _extract_job_title(job),
        },
    }


@app.delete("/api/admin/jobs/{job_id}")
async def admin_delete_job(job_id: str, request: Request):
    _require_admin(request)
    job = _get_generation_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="任务不存在")
    _delete_generation_job(job_id)
    with _PROGRESS_LOCK:
        _PROGRESS_EVENTS.pop(job_id, None)
        _PROGRESS_DONE.pop(job_id, None)
    return {"status": "success", "job_id": job_id}


@app.post("/api/admin/users")
async def admin_create_user(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    role: str = Form("user"),
):
    _require_admin(request)
    try:
        created = _create_user(username, password, role=role)
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=409, detail="账号已存在")
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return {"status": "success", "user": created}


@app.post("/api/admin/users/batch")
async def admin_batch_create_users(
    request: Request,
    prefix: str = Form("tester"),
    count: int = Form(5),
    role: str = Form("user"),
):
    _require_admin(request)
    try:
        users = _batch_create_users(prefix=prefix, count=count, role=role)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return {"status": "success", "users": users}


def _generate_whitepaper_sync(data_text: str) -> str:
    research_context = ""
    references: List[Dict[str, str]] = []
    if ENABLE_WEB_SEARCH:
        try:
            _progress_print("[Research] 开始联网检索与资料筛选...", stage="research")
            research_topic = _extract_research_topic_with_ai(data_text)
            _progress_print(f"[Research][Topic] AI拆词: {research_topic}", stage="research")
            research_debug = collect_recent_references_debug(research_topic, limit=10, timeout=WEB_SEARCH_TIMEOUT)
            references = research_debug.get("final_results", [])
            research_context = build_research_context(references)
            _progress_print(
                "[Research] 可用资料: "
                f"{len(references)} 条 / 候选: {research_debug.get('candidate_count', 0)} / "
                f"已抓取: {research_debug.get('enriched_count', 0)} / "
                f"淘汰: {len(research_debug.get('rejected_results', []))}",
                stage="research",
            )
            for log in research_debug.get("search_logs", [])[:20]:
                if log.get("error"):
                    _progress_print(
                        f"[Research][Search] {log.get('source')} query={log.get('query')} "
                        f"失败: {log.get('error')}",
                        stage="research",
                    )
                else:
                    _progress_print(
                        f"[Research][Search] {log.get('source')} query={log.get('query')} "
                        f"results={log.get('result_count')}",
                        stage="research",
                    )
            for log in research_debug.get("fetch_logs", [])[:30]:
                _progress_print(
                    f"[Research][Fetch] score={log.get('score')} status={log.get('status')} "
                    f"source={log.get('publisher')} title={log.get('title')} url={log.get('url')}",
                    stage="research",
                )
            for idx, ref in enumerate(references, 1):
                _progress_print(
                    f"[Research] 入选 {idx}: score={ref.get('match_score')} "
                    f"status={ref.get('access_status')} source={ref.get('publisher')} "
                    f"title={ref.get('title')}",
                    stage="research",
                )
            for ref in research_debug.get("rejected_results", [])[:8]:
                _progress_print(
                    f"[Research] 淘汰: reason={ref.get('reject_reason')} "
                    f"score={ref.get('match_score')} status={ref.get('access_status')} "
                    f"source={ref.get('publisher')} title={ref.get('title')}",
                    stage="research",
                )
        except Exception as e:
            _progress_print(f"[Research] 联网检索失败: {e}", stage="research")

    try:
        _progress_print("[Pipeline] 开始四角色白皮书生成...", stage="pipeline")
        return run_four_agent_pipeline(
            data_text,
            research_context=research_context,
            references=references,
        )
    except (RuntimeError, ValueError):
        import traceback
        traceback.print_exc()
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise RuntimeError(f"AI 生成失败: {str(e)}") from e


@app.post("/api/generate")
async def generate_whitepaper(
    request: Request,
    files: List[UploadFile] = File([]),
    text: Optional[str] = Form(None),
    job_id: Optional[str] = Form(None),
):
    """
    主接口：四角色管线 — 战略官 → 主笔+视觉导演 → 图表数据官
    支持同时上传多个文件 + 可选文字输入
    """
    # 1. 获取数据（多文件 + 文字拼合）
    data_parts = []
    for f in files:
        if not f.filename:
            continue
        content = await f.read()
        if not content:
            raise HTTPException(status_code=400, detail=f"文件 {f.filename} 为空")
        try:
            parsed = FileParser.parse(f.filename, content)
            data_parts.append(f"【文件: {f.filename}】\n{parsed}")
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
    if text and text.strip():
        data_parts.append(f"【文字输入】\n{text.strip()}")
    if not data_parts:
        raise HTTPException(status_code=400, detail="请上传文件或输入文字")
    data_text = "\n\n---\n\n".join(data_parts)

    # 2. 截断过长内容
    MAX_CHARS = 15000
    if len(data_text) > MAX_CHARS:
        data_text = data_text[:MAX_CHARS] + "\n\n...(数据已截断，共{}字符)".format(len(data_text))

    active_job_id = (job_id or f"job-{uuid.uuid4().hex}").strip()
    with _PROGRESS_LOCK:
        _PROGRESS_EVENTS.setdefault(active_job_id, [])
        _PROGRESS_DONE[active_job_id] = False

    try:
        _init_app_db()
        created_job_id = _create_generation_job(_current_username(request), data_text, active_job_id)
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=409, detail="任务 ID 已存在，请重新提交")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"任务创建失败: {str(e)}")

    _emit_progress("[Input] 数据解析完成，任务已进入后台队列", job_id=created_job_id, stage="input")
    _enqueue_generation_job(created_job_id)
    return {"status": "queued", "job_id": created_job_id}


@app.post("/api/generate-image")
async def generate_image(
    prompt: str = Form(...),
    width: int = Form(768),
    height: int = Form(512),
):
    """
    配图生成接口：前端解析到 <image> 标签后调用此接口
    返回 base64 图片或占位提示
    """
    if not image_generator.available:
        return {
            "status": "unavailable",
            "message": "豆包星绘未配置，请在 .env 中设置 VOLCENGINE_ACCESS_KEY 和 VOLCENGINE_SECRET_KEY",
            "image_base64": None,
        }

    # Agent3 已生成高质量中文提示词，直接使用
    image_b64 = image_generator.generate(prompt, width=width, height=height)

    if image_b64:
        return {
            "status": "success",
            "image_base64": image_b64,
        }
    else:
        return {
            "status": "error",
            "message": "图片生成失败，请稍后重试",
            "image_base64": None,
        }


@app.post("/api/research-sources")
async def research_sources(
    text: str = Form(...),
    limit: int = Form(10),
    debug: bool = Form(False),
):
    """
    轻量联网检索接口：仅返回外部来源结果，不触发整篇白皮书生成。
    用于验证联网搜索是否正常返回信息。
    """
    query = (text or "").strip()
    if not query:
        raise HTTPException(status_code=400, detail="请输入检索主题")

    capped_limit = max(1, min(limit, 10))

    try:
        if debug:
            debug_payload = collect_recent_references_debug(
                query,
                limit=capped_limit,
                timeout=WEB_SEARCH_TIMEOUT,
            )
            results = debug_payload.get("final_results", [])
            return {
                "status": "success",
                "query": query,
                "count": len(results),
                "results": results,
                "debug": {
                    key: value
                    for key, value in debug_payload.items()
                    if key != "final_results"
                },
                "timestamp": int(time.time()),
            }

        results = collect_recent_references(query, limit=capped_limit, timeout=WEB_SEARCH_TIMEOUT)
        return {
            "status": "success",
            "query": query,
            "count": len(results),
            "results": results,
            "timestamp": int(time.time()),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"联网检索失败: {str(e)}")


@app.post("/api/verify-data")
async def verify_data(
    files: List[UploadFile] = File([]),
    text: Optional[str] = Form(None),
):
    """
    数据验证接口：验证上传的数据文件质量和结构
    支持多文件 + 可选文字输入的验证
    """
    if not files and not text:
        raise HTTPException(status_code=400, detail="请上传文件或输入文字进行验证")
    
    verification_results = []
    total_files = len(files)
    
    # 验证文件
    for i, file in enumerate(files):
        if not file.filename:
            continue
            
        try:
            content = await file.read()
            if not content:
                verification_results.append({
                    "filename": file.filename,
                    "status": "error",
                    "message": "文件内容为空",
                    "details": {}
                })
                continue
                
            # 检查文件格式
            ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
            supported_formats = ["csv", "xlsx", "xls", "docx", "pdf"]
            
            if ext not in supported_formats:
                verification_results.append({
                    "filename": file.filename,
                    "status": "error",
                    "message": f"不支持的文件格式: .{ext}",
                    "details": {
                        "supported_formats": supported_formats,
                        "detected_format": ext
                    }
                })
                continue
            
            # 解析文件内容进行验证
            try:
                parsed_content = FileParser.parse(file.filename, content)
                
                # 提取结构化信息用于验证
                if ext in ["csv", "xlsx", "xls"]:
                    # 检查CSV/Excel数据质量
                    df = pd.read_excel(io.BytesIO(content)) if ext in ["xlsx", "xls"] else pd.read_csv(io.BytesIO(content), encoding="utf-8")
                    
                    quality_check = {
                        "file_type": "spreadsheet",
                        "rows": len(df),
                        "columns": len(df.columns),
                        "column_names": df.columns.tolist(),
                        "missing_values": df.isnull().sum().to_dict(),
                        "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
                        "numeric_columns": df.select_dtypes(include=['number']).columns.tolist(),
                        "text_columns": df.select_dtypes(include=['object']).columns.tolist(),
                        "sample_data": df.head(3).to_dict('records')
                    }
                    
                    # 数据质量评分
                    missing_percentage = (df.isnull().sum().sum() / (len(df) * len(df.columns))) * 100
                    quality_score = max(0, 100 - missing_percentage)
                    
                    verification_results.append({
                        "filename": file.filename,
                        "status": "success" if quality_score > 50 else "warning",
                        "message": f"数据质量评分: {quality_score:.1f}/100",
                        "details": {
                            "quality_check": quality_check,
                            "quality_score": quality_score,
                            "recommendations": _generate_data_recommendations(df, quality_score)
                        }
                    })
                    
                elif ext == "docx":
                    # 检查Word文档质量
                    doc = Document(io.BytesIO(content))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    
                    doc_check = {
                        "file_type": "document",
                        "paragraphs": len(paragraphs),
                        "word_count": sum(len(p.split()) for p in paragraphs),
                        "has_content": len(paragraphs) > 0,
                        "sample_paragraphs": paragraphs[:3]
                    }
                    
                    verification_results.append({
                        "filename": file.filename,
                        "status": "success" if len(paragraphs) > 5 else "warning",
                        "message": f"文档包含 {len(paragraphs)} 个段落",
                        "details": doc_check
                    })
                    
                elif ext == "pdf":
                    # 检查PDF文档质量
                    with pdfplumber.open(io.BytesIO(content)) as pdf:
                        pages = pdf.pages
                        text_content = []
                        for page in pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_content.append(page_text)
                        
                        pdf_check = {
                            "file_type": "pdf",
                            "pages": len(pages),
                            "has_text": len(text_content) > 0,
                            "total_text_length": sum(len(text) for text in text_content),
                            "sample_text": text_content[0][:500] if text_content else ""
                        }
                        
                        verification_results.append({
                            "filename": file.filename,
                            "status": "success" if len(text_content) > 0 else "warning",
                            "message": f"PDF包含 {len(pages)} 页，提取到 {len(text_content)} 页文本",
                            "details": pdf_check
                        })
                        
            except Exception as parse_error:
                verification_results.append({
                    "filename": file.filename,
                    "status": "error",
                    "message": f"文件解析失败: {str(parse_error)}",
                    "details": {}
                })
                
        except Exception as e:
            verification_results.append({
                "filename": file.filename or f"file_{i+1}",
                "status": "error",
                "message": f"验证过程中发生错误: {str(e)}",
                "details": {}
            })
    
    # 验证文本输入
    if text and text.strip():
        text_check = {
            "type": "text_input",
            "length": len(text),
            "word_count": len(text.split()),
            "has_content": len(text.strip()) > 0,
            "sample": text[:200] + "..." if len(text) > 200 else text
        }
        
        verification_results.append({
            "filename": "text_input",
            "status": "success",
            "message": f"文本输入有效，包含 {len(text.split())} 个单词",
            "details": text_check
        })
    
    # 汇总结果
    total_items = len(verification_results)
    success_count = sum(1 for r in verification_results if r["status"] == "success")
    warning_count = sum(1 for r in verification_results if r["status"] == "warning")
    error_count = sum(1 for r in verification_results if r["status"] == "error")
    
    summary = {
        "total_items": total_items,
        "success": success_count,
        "warnings": warning_count,
        "errors": error_count,
        "overall_status": "success" if error_count == 0 else "warning" if error_count < total_items else "error"
    }
    
    return {
        "status": "completed",
        "summary": summary,
        "results": verification_results,
        "timestamp": int(time.time())
    }


def _generate_data_recommendations(df: pd.DataFrame, quality_score: float) -> List[str]:
    """根据数据质量生成改进建议"""
    recommendations = []
    
    # 检查缺失值
    missing_percentage = (df.isnull().sum().sum() / (len(df) * len(df.columns))) * 100
    if missing_percentage > 20:
        recommendations.append(f"数据缺失率较高({missing_percentage:.1f}%)，建议补充数据或删除过多缺失的列")
    elif missing_percentage > 5:
        recommendations.append(f"存在少量缺失数据({missing_percentage:.1f}%)，建议考虑插值或删除")
    
    # 检查列数量
    if len(df.columns) < 3:
        recommendations.append("建议增加更多数据列以提高分析维度")
    elif len(df.columns) > 20:
        recommendations.append("列数较多，建议筛选关键列或进行特征工程")
    
    # 检查数值列
    numeric_cols = df.select_dtypes(include=['number']).columns
    if len(numeric_cols) == 0:
        recommendations.append("建议包含数值型列以支持量化分析")
    
    # 检查数据量
    if len(df) < 10:
        recommendations.append("数据量较少，建议增加样本数量")
    elif len(df) > 1000:
        recommendations.append("数据量较大，可考虑抽样以提高处理效率")
    
    # 如果质量很好，给出肯定
    if quality_score > 90:
        recommendations.append("数据质量优秀，适合直接用于分析")
    
    return recommendations if recommendations else ["数据格式正确，可直接使用"]


@app.post("/api/reflect-report")
async def create_reflection_report(
    report_id: str = Form(...),
    content: str = Form(...),
    reflection_type: str = Form("general"),
    source_data: Optional[str] = Form(None),
    analysis_metadata: Optional[str] = Form(None),
):
    """
    创建反思报告接口
    接收反思报告内容并存储，支持报告ID、类型和元数据
    """
    try:
        # 解析元数据（如果提供）
        metadata = {
            "type": reflection_type,
            "source_data": source_data,
            "analysis_metadata": json.loads(analysis_metadata) if analysis_metadata else {},
            "created_by": "system"
        }
        
        # 存储反思报告
        reflection = reflection_storage.add_reflection(
            report_id=report_id,
            content=content,
            metadata=metadata
        )
        
        return {
            "status": "success",
            "message": "反思报告创建成功",
            "reflection_id": report_id,
            "timestamp": reflection["timestamp"],
            "metadata": metadata
        }
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="analysis_metadata 必须是有效的 JSON 字符串")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"创建反思报告失败: {str(e)}")


@app.get("/api/reflection-history")
async def get_reflection_history(
    limit: int = Query(50, ge=1, le=200, description="返回记录数量限制"),
    offset: int = Query(0, ge=0, description="偏移量"),
    reflection_type: Optional[str] = Query(None, description="按类型筛选"),
):
    """
    获取反思历史接口
    支持分页和类型筛选，返回反思报告的历史记录
    """
    try:
        if reflection_type:
            # 按类型筛选
            reflections = reflection_storage.get_reflections_by_type(reflection_type, limit)
            total_count = len(reflections)
        else:
            # 获取全部历史
            reflections = reflection_storage.get_reflection_history(limit, offset)
            total_count = len(reflection_storage.reflections)
        
        return {
            "status": "success",
            "reflections": reflections,
            "pagination": {
                "limit": limit,
                "offset": offset,
                "total": total_count,
                "has_more": offset + limit < total_count
            },
            "timestamp": int(time.time())
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取反思历史失败: {str(e)}")


@app.get("/api/reflect-report/{report_id}")
async def get_reflection_report(report_id: str):
    """
    获取指定反思报告接口
    根据报告ID获取单个反思报告的详细信息
    """
    try:
        reflection = reflection_storage.get_reflection(report_id)
        
        if not reflection:
            raise HTTPException(status_code=404, detail=f"未找到反思报告: {report_id}")
        
        return {
            "status": "success",
            "reflection": reflection,
            "timestamp": int(time.time())
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取反思报告失败: {str(e)}")


@app.delete("/api/reflect-report/{report_id}")
async def delete_reflection_report(report_id: str):
    """
    删除反思报告接口
    根据报告ID删除指定的反思报告
    """
    try:
        reflection = reflection_storage.get_reflection(report_id)
        
        if not reflection:
            raise HTTPException(status_code=404, detail=f"未找到反思报告: {report_id}")
        
        # 从存储中删除
        with reflection_storage.lock:
            reflection_storage.reflections = [
                r for r in reflection_storage.reflections 
                if r["id"] != report_id
            ]
        
        return {
            "status": "success",
            "message": f"反思报告 {report_id} 已删除",
            "timestamp": int(time.time())
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除反思报告失败: {str(e)}")


# ---------------------------------------------------------------------------
# 启动入口
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys
    # Windows 终端编码兼容
    if sys.platform == "win32":
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print("=" * 50)
    print("  CIBE 美业白皮书生成系统 启动中...")
    print(f"  DeepSeek API: {'已配置' if DEEPSEEK_API_KEY else '未配置'}")
    print(f"  豆包星绘:     {'已配置' if image_generator.available else '未配置(配图功能不可用)'}")
    print(f"  访问地址: http://localhost:5678")
    print("=" * 50)
    uvicorn.run(app, host="0.0.0.0", port=5678)
